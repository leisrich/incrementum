# core/document_processor/handlers/html_handler.py

import os
import logging
import tempfile
import requests
from typing import Dict, Any, Optional, List, Tuple
from urllib.parse import urlparse
from datetime import datetime
from docx import Document

from bs4 import BeautifulSoup

from .base_handler import DocumentHandler

logger = logging.getLogger(__name__)

class HTMLHandler(DocumentHandler):
    """Handler for processing HTML documents."""
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from an HTML document."""
        metadata = {
            'title': os.path.basename(file_path),
            'author': '',
            'creation_date': None,
            'modification_date': None
        }
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f.read(), 'lxml')
                
                # Extract title
                if soup.title:
                    metadata['title'] = soup.title.string
                
                # Extract author
                author_meta = soup.find('meta', attrs={'name': 'author'})
                if author_meta and author_meta.get('content'):
                    metadata['author'] = author_meta['content']
                
                # Extract other metadata
                description_meta = soup.find('meta', attrs={'name': 'description'})
                if description_meta and description_meta.get('content'):
                    metadata['description'] = description_meta['content']
                
        except Exception as e:
            logger.exception(f"Error extracting HTML metadata: {e}")
        
        return metadata
    
    def extract_content(self, file_path: str) -> Dict[str, Any]:
        """Extract content from an HTML document."""
        result = {
            'text': '',
            'html': '',
            'elements': []
        }
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
                
                soup = BeautifulSoup(html_content, 'lxml')
                
                # Extract text content
                result['text'] = soup.get_text(separator='\n')
                
                # Store original HTML
                result['html'] = html_content
                
                # Extract elements
                for element in soup.find_all(['h1', 'h2', 'h3', 'p']):
                    result['elements'].append({
                        'type': element.name,
                        'content': element.get_text(),
                        'html': str(element)
                    })
                
        except Exception as e:
            logger.exception(f"Error extracting HTML content: {e}")
        
        return result
    
    def download_from_url(self, url: str) -> Tuple[Optional[str], Dict[str, Any]]:
        """
        Download an HTML document from a URL and save it to permanent storage.
        
        Args:
            url: URL to download from
            
        Returns:
            Tuple of (local file path, metadata)
        """
        metadata = {}
        
        try:
            # Import document storage directory
            from core.document_processor.document_importer import DOCUMENT_STORAGE_DIR
            
            # Generate a unique filename based on URL and timestamp
            filename = f"html_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{os.path.basename(url)}"
            if not filename.endswith('.html'):
                filename += '.html'
            
            # Sanitize filename
            filename = ''.join(c for c in filename if c.isalnum() or c in '._-')
            
            # Full path to save the file
            file_path = os.path.join(DOCUMENT_STORAGE_DIR, filename)
            
            # Download the file
            response = requests.get(url, stream=True)
            response.raise_for_status()
            
            # Save the content
            with open(file_path, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    f.write(chunk)
            
            # Extract metadata
            metadata = self.extract_metadata(file_path)
            metadata['source_url'] = url
            
            return file_path, metadata
            
        except Exception as e:
            logger.exception(f"Error downloading HTML: {e}")
            return None, metadata


# core/document_processor/handlers/text_handler.py

import os
import logging
import tempfile
import requests
from typing import Dict, Any, Optional, List, Tuple
from datetime import datetime

from .base_handler import DocumentHandler

logger = logging.getLogger(__name__)

class TextHandler(DocumentHandler):
    """Handler for processing plain text documents."""
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from a text document."""
        metadata = {
            'title': os.path.basename(file_path),
            'author': '',
            'creation_date': None,
            'modification_date': None
        }
        
        # For plain text, we don't have much metadata
        # Just use file stats
        try:
            stat = os.stat(file_path)
            metadata['creation_date'] = datetime.fromtimestamp(stat.st_ctime)
            metadata['modification_date'] = datetime.fromtimestamp(stat.st_mtime)
            
        except Exception as e:
            logger.exception(f"Error extracting text metadata: {e}")
        
        return metadata
    
    def extract_content(self, file_path: str) -> Dict[str, Any]:
        """Extract content from a text document."""
        result = {
            'text': '',
            'lines': [],
            'paragraphs': []
        }
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                
                # Store full text
                result['text'] = content
                
                # Split into lines
                result['lines'] = content.splitlines()
                
                # Split into paragraphs (separated by blank lines)
                result['paragraphs'] = [p for p in content.split('\n\n') if p.strip()]
                
        except Exception as e:
            logger.exception(f"Error extracting text content: {e}")
        
        return result
    
    def download_from_url(self, url: str) -> Tuple[Optional[str], Dict[str, Any]]:
        """Download a text document from a URL."""
        metadata = {}
        
        try:
            # Download the file
            response = requests.get(url, stream=True)
            response.raise_for_status()
            
            # Create a temporary file
            fd, temp_path = tempfile.mkstemp(suffix='.txt')
            
            # Save the content
            with os.fdopen(fd, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    f.write(chunk)
            
            # Extract metadata
            metadata = self.extract_metadata(temp_path)
            metadata['source_url'] = url
            
            return temp_path, metadata
            
        except Exception as e:
            logger.exception(f"Error downloading text: {e}")
            return None, metadata


# core/document_processor/handlers/epub_handler.py

import os
import logging
import tempfile
import requests
from typing import Dict, Any, Optional, List, Tuple
from datetime import datetime

import ebooklib
from ebooklib import epub
from bs4 import BeautifulSoup

from .base_handler import DocumentHandler

logger = logging.getLogger(__name__)

class EPUBHandler(DocumentHandler):
    """Handler for processing EPUB documents."""
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from an EPUB document."""
        metadata = {
            'title': os.path.basename(file_path),
            'author': '',
            'creation_date': None,
            'modification_date': None,
            'language': '',
            'publisher': '',
            'toc': []
        }
        
        try:
            book = epub.read_epub(file_path)
            
            # Extract metadata
            if book.get_metadata('DC', 'title'):
                metadata['title'] = book.get_metadata('DC', 'title')[0][0]
            
            if book.get_metadata('DC', 'creator'):
                metadata['author'] = book.get_metadata('DC', 'creator')[0][0]
            
            if book.get_metadata('DC', 'date'):
                date_str = book.get_metadata('DC', 'date')[0][0]
                try:
                    metadata['creation_date'] = datetime.fromisoformat(date_str)
                except ValueError:
                    pass
            
            if book.get_metadata('DC', 'language'):
                metadata['language'] = book.get_metadata('DC', 'language')[0][0]
            
            if book.get_metadata('DC', 'publisher'):
                metadata['publisher'] = book.get_metadata('DC', 'publisher')[0][0]
            
            # Extract table of contents
            toc = book.toc
            for item in toc:
                if isinstance(item, tuple):
                    metadata['toc'].append({
                        'title': item[0].title,
                        'href': item[0].href
                    })
                else:
                    metadata['toc'].append({
                        'title': item.title,
                        'href': item.href
                    })
            
        except Exception as e:
            logger.exception(f"Error extracting EPUB metadata: {e}")
        
        return metadata
    
    def extract_content(self, file_path: str) -> Dict[str, Any]:
        """Extract content from an EPUB document."""
        result = {
            'text': '',
            'chapters': [],
            'toc': []
        }
        
        try:
            book = epub.read_epub(file_path)
            
            # Extract table of contents
            toc = book.toc
            for item in toc:
                if isinstance(item, tuple):
                    result['toc'].append({
                        'title': item[0].title,
                        'href': item[0].href
                    })
                else:
                    result['toc'].append({
                        'title': item.title,
                        'href': item.href
                    })
            
            # Extract content
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    try:
                        content = item.get_content().decode('utf-8')
                        soup = BeautifulSoup(content, 'lxml')
                        
                        chapter_text = soup.get_text(separator='\n')
                        result['text'] += chapter_text + '\n\n'
                        
                        # Add chapter
                        result['chapters'].append({
                            'id': item.id,
                            'href': item.get_name(),
                            'content': chapter_text
                        })
                    except Exception as e:
                        logger.warning(f"Error processing EPUB item {item.id}: {e}")
            
        except Exception as e:
            logger.exception(f"Error extracting EPUB content: {e}")
        
        return result
    
    def download_from_url(self, url: str) -> Tuple[Optional[str], Dict[str, Any]]:
        """Download an EPUB document from a URL."""
        metadata = {}
        
        try:
            # Download the file
            response = requests.get(url, stream=True)
            response.raise_for_status()
            
            # Create a temporary file
            fd, temp_path = tempfile.mkstemp(suffix='.epub')
            
            # Save the content
            with os.fdopen(fd, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    f.write(chunk)
            
            # Extract metadata
            metadata = self.extract_metadata(temp_path)
            metadata['source_url'] = url
            
            return temp_path, metadata
            
        except Exception as e:
            logger.exception(f"Error downloading EPUB: {e}")
            return None, metadata


# core/document_processor/handlers/docx_handler.py

import os
import logging
import tempfile
import requests
from typing import Dict, Any, Optional, List, Tuple
from datetime import datetime

import docx

from .base_handler import DocumentHandler

logger = logging.getLogger(__name__)

class DOCXHandler(DocumentHandler):
    """Handler for processing DOCX documents."""
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from a DOCX document."""
        metadata = {
            'title': os.path.basename(file_path),
            'author': '',
            'creation_date': None,
            'modification_date': None
        }
        
        try:
            doc = docx.Document(file_path)
            
            # Extract core properties
            if doc.core_properties.title:
                metadata['title'] = doc.core_properties.title
            
            if doc.core_properties.author:
                metadata['author'] = doc.core_properties.author
            
            if doc.core_properties.created:
                metadata['creation_date'] = doc.core_properties.created
            
            if doc.core_properties.modified:
                metadata['modification_date'] = doc.core_properties.modified
            
        except Exception as e:
            logger.exception(f"Error extracting DOCX metadata: {e}")
        
        return metadata
    
    def extract_content(self, file_path: str) -> Dict[str, Any]:
        """Extract content from a DOCX document."""
        result = {
            'text': '',
            'paragraphs': [],
            'headings': []
        }
        
        try:
            doc = docx.Document(file_path)
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    result['text'] += text + '\n\n'
                    
                    # Check if it's a heading
                    if paragraph.style.name.startswith('Heading'):
                        result['headings'].append({
                            'level': int(paragraph.style.name.replace('Heading', '')),
                            'text': text
                        })
                    else:
                        result['paragraphs'].append(text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = '\t'.join(cell.text for cell in row.cells)
                    result['text'] += row_text + '\n'
            
        except Exception as e:
            logger.exception(f"Error extracting DOCX content: {e}")
        
        return result
    
    def download_from_url(self, url: str) -> Tuple[Optional[str], Dict[str, Any]]:
        """Download a DOCX document from a URL."""
        metadata = {}
        
        try:
            # Download the file
            response = requests.get(url, stream=True)
            response.raise_for_status()
            
            # Create a temporary file
            fd, temp_path = tempfile.mkstemp(suffix='.docx')
            
            # Save the content
            with os.fdopen(fd, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    f.write(chunk)
            
            # Extract metadata
            metadata = self.extract_metadata(temp_path)
            metadata['source_url'] = url
            
            return temp_path, metadata
            
        except Exception as e:
            logger.exception(f"Error downloading DOCX: {e}")
            return None, metadata
