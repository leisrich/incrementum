# core/document_processor/handlers/docx_handler.py

import os
import logging
import tempfile
import requests
from typing import Dict, Any, Optional, List, Tuple
from datetime import datetime

from docx import Document as DocxDocument
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI

from .base_handler import DocumentHandler

logger = logging.getLogger(__name__)

class DOCXHandler(DocumentHandler):
    """Handler for processing DOCX documents."""
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from a DOCX document.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dictionary of metadata
        """
        metadata = {
            'title': os.path.basename(file_path),
            'author': '',
            'creation_date': None,
            'modification_date': None,
            'comments': '',
            'categories': '',
            'subject': '',
            'keywords': '',
            'version': '',
            'content_status': '',
            'language': '',
            'word_count': 0,
            'paragraph_count': 0
        }
        
        try:
            # Load document
            doc = DocxDocument(file_path)
            
            # Extract core properties
            core_props = doc.core_properties
            
            if core_props.title:
                metadata['title'] = core_props.title
            
            if core_props.author:
                metadata['author'] = core_props.author
            
            if core_props.created:
                metadata['creation_date'] = core_props.created
            
            if core_props.modified:
                metadata['modification_date'] = core_props.modified
            
            if core_props.comments:
                metadata['comments'] = core_props.comments
            
            if core_props.category:
                metadata['categories'] = core_props.category
            
            if core_props.subject:
                metadata['subject'] = core_props.subject
            
            if core_props.keywords:
                metadata['keywords'] = core_props.keywords
            
            if core_props.version:
                metadata['version'] = core_props.version
            
            if core_props.content_status:
                metadata['content_status'] = core_props.content_status
            
            if core_props.language:
                metadata['language'] = core_props.language
            
            # Count paragraphs
            metadata['paragraph_count'] = len(doc.paragraphs)
            
            # Estimate word count
            word_count = 0
            for para in doc.paragraphs:
                word_count += len(para.text.split())
            metadata['word_count'] = word_count
            
        except Exception as e:
            logger.exception(f"Error extracting DOCX metadata: {e}")
        
        return metadata
    
    def extract_content(self, file_path: str) -> Dict[str, Any]:
        """
        Extract content from a DOCX document.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dictionary containing extracted content and structure
        """
        result = {
            'text': '',
            'paragraphs': [],
            'headings': [],
            'tables': [],
            'sections': [],
            'structure': []
        }
        
        try:
            # Load document
            doc = DocxDocument(file_path)
            
            # Process document structure
            current_section = {'title': '', 'content': '', 'level': 0, 'children': []}
            current_heading = None
            
            # Process paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                
                # Skip empty paragraphs
                if not text:
                    continue
                
                # Add to full text
                result['text'] += text + '\n\n'
                
                # Check if it's a heading
                is_heading = False
                heading_level = 0
                
                if paragraph.style and paragraph.style.name:
                    if paragraph.style.name.startswith('Heading'):
                        try:
                            heading_level = int(paragraph.style.name.replace('Heading', ''))
                            is_heading = True
                        except ValueError:
                            # Not a standard heading style
                            pass
                
                if is_heading:
                    # Add to headings
                    heading = {
                        'text': text,
                        'level': heading_level
                    }
                    result['headings'].append(heading)
                    
                    # Update structure
                    if heading_level == 1:
                        # Top-level heading
                        if current_section['title']:
                            # Add previous section
                            result['sections'].append(current_section)
                        
                        # Start new section
                        current_section = {
                            'title': text,
                            'content': '',
                            'level': heading_level,
                            'children': []
                        }
                    else:
                        # Sub-heading
                        current_section['children'].append({
                            'title': text,
                            'content': '',
                            'level': heading_level,
                            'children': []
                        })
                    
                    current_heading = {'text': text, 'level': heading_level}
                else:
                    # Regular paragraph
                    para_dict = {
                        'text': text,
                        'heading': current_heading.copy() if current_heading else None
                    }
                    result['paragraphs'].append(para_dict)
                    
                    # Add to current section
                    if current_section['title']:
                        current_section['content'] += text + '\n\n'
            
            # Add final section
            if current_section['title']:
                result['sections'].append(current_section)
            
            # Process tables
            for i, table in enumerate(doc.tables):
                table_data = {
                    'index': i,
                    'rows': len(table.rows),
                    'columns': len(table.columns) if table.rows else 0,
                    'data': []
                }
                
                # Extract table data
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    
                    table_data['data'].append(row_data)
                
                result['tables'].append(table_data)
                
                # Add table to text
                result['text'] += f"\nTable {i+1}:\n"
                for row in table_data['data']:
                    result['text'] += ' | '.join(row) + '\n'
                result['text'] += '\n'
            
            # Create document structure
            result['structure'] = self._extract_document_structure(doc)
            
        except Exception as e:
            logger.exception(f"Error extracting DOCX content: {e}")
        
        return result
    
    def _extract_document_structure(self, doc) -> List[Dict[str, Any]]:
        """
        Extract hierarchical document structure.
        
        Args:
            doc: The DOCX document
            
        Returns:
            List of dictionaries representing document structure
        """
        structure = []
        current_level = [structure]
        current_heading_level = 0
        
        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                continue
                
            is_heading = False
            heading_level = 0
            
            if paragraph.style and paragraph.style.name:
                if paragraph.style.name.startswith('Heading'):
                    try:
                        heading_level = int(paragraph.style.name.replace('Heading', ''))
                        is_heading = True
                    except ValueError:
                        pass
            
            if is_heading:
                # Create new heading node
                heading_node = {
                    'type': 'heading',
                    'text': paragraph.text,
                    'level': heading_level,
                    'children': []
                }
                
                # Find correct level to insert
                if heading_level > current_heading_level:
                    # Go deeper
                    if current_level[-1]:
                        parent = current_level[-1][-1]
                        current_level.append(parent['children'])
                elif heading_level < current_heading_level:
                    # Go up
                    levels_up = current_heading_level - heading_level
                    for _ in range(levels_up + 1):
                        if len(current_level) > 1:
                            current_level.pop()
                
                # Add to current level
                current_level[-1].append(heading_node)
                current_heading_level = heading_level
            else:
                # Create paragraph node
                para_node = {
                    'type': 'paragraph',
                    'text': paragraph.text
                }
                
                # Add to current level
                if current_level[-1]:
                    # If we have a heading, add as child
                    if current_level[-1][-1]['type'] == 'heading':
                        current_level[-1][-1]['children'].append(para_node)
                    else:
                        # Add at current level
                        current_level[-1].append(para_node)
                else:
                    # No structure yet, add to root
                    current_level[0].append(para_node)
        
        return structure
    
    def download_from_url(self, url: str) -> Tuple[Optional[str], Dict[str, Any]]:
        """
        Download a DOCX document from a URL.
        
        Args:
            url: URL to download from
            
        Returns:
            Tuple of (local file path, metadata)
        """
        metadata = {}
        
        try:
            # Download the file
            response = requests.get(url, stream=True)
            response.raise_for_status()
            
            # Create a temporary file
            fd, temp_path = tempfile.mkstemp(suffix='.docx')
            
            # Save the content
            with os.fdopen(fd, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    f.write(chunk)
            
            # Extract metadata
            metadata = self.extract_metadata(temp_path)
            metadata['source_url'] = url
            
            return temp_path, metadata
            
        except Exception as e:
            logger.exception(f"Error downloading DOCX: {e}")
            return None, metadata
