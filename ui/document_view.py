# If the patching approach doesn't work, here's a complete replacement for document_view.py
# Save this as ui/document_view.py.new and rename it if needed

import os
import logging
from typing import Optional, List, Dict
from datetime import datetime
import json
import time
import tempfile
from io import BytesIO
from pathlib import Path
import base64

from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, 
    QPushButton, QScrollArea, QSplitter, QTextEdit,
    QToolBar, QMenu, QMessageBox, QApplication, QDialog,
    QTabWidget, QLineEdit, QSizePolicy, QCheckBox, QSlider
)
from PyQt6.QtCore import Qt, pyqtSignal, pyqtSlot, QPoint, QUrl, QObject, QTimer, QPointF, QSize, QByteArray
from PyQt6.QtGui import QAction, QTextCursor, QColor, QTextCharFormat, QKeyEvent, QIntValidator
try:
    from PyQt6.QtWebEngineWidgets import QWebEngineView
    from PyQt6.QtWebEngineCore import QWebEngineSettings
    from PyQt6.QtWebChannel import QWebChannel
    HAS_WEBENGINE = True
except ImportError:
    HAS_WEBENGINE = False

from core.knowledge_base.models import Document, Extract
from core.content_extractor.extractor import ContentExtractor
from core.document_processor.handlers.epub_handler import EPUBHandler
from .document_extracts_view import DocumentExtractsView
from .load_epub_helper import setup_epub_webview
from .load_youtube_helper import setup_youtube_webview, extract_video_id_from_document
from .youtube_transcript_view import YouTubeTranscriptView

logger = logging.getLogger(__name__)

class WebViewCallback(QObject):
    """Callback handler for JavaScript communication with WebView."""
    
    def __init__(self, document_view):
        """Initialize with document view reference."""
        super().__init__()
        self.document_view = document_view
    
    @pyqtSlot(str)
    def selectionChanged(self, text):
        """Handle text selection changes from WebView."""
        if self.document_view:
            self.document_view._handle_webview_selection(text)
    
    @pyqtSlot(str)
    def extractText(self, text):
        """Handle text extraction requests from WebView."""
        if self.document_view and text:
            self.document_view._handle_sm_extract_result(text)
    
    @pyqtSlot(str)
    def createCloze(self, text):
        """Handle cloze creation requests from WebView."""
        if self.document_view and text:
            self.document_view._handle_sm_cloze_result(text)
    
    @pyqtSlot()
    def skipItem(self):
        """Handle skip item requests from WebView."""
        if self.document_view:
            # This would connect to review scheduling
            pass

class VimKeyHandler:
    """Helper class for handling Vim-like key bindings."""
    
    def __init__(self, document_view):
        self.document_view = document_view
        self.vim_mode = True  # Default to Vim mode on
        self.command_mode = False  # Normal mode by default (not command mode)
        self.visual_mode = False  # Not in visual mode by default
        self.current_command = ""
        self.count_prefix = ""  # For number prefixes like 5j
        self.selection_start = None  # For visual mode selection
        self.selection_active = False
        
    def toggle_vim_mode(self):
        """Toggle Vim mode on/off."""
        self.vim_mode = not self.vim_mode
        if not self.vim_mode:
            self.command_mode = False
            self.visual_mode = False
            self.current_command = ""
        logger.debug(f"Vim mode {'enabled' if self.vim_mode else 'disabled'}")
        return self.vim_mode
        
    def handle_key_event(self, event):
        """Handle key events in Vim style."""
        if not self.vim_mode:
            return False
            
        # Get key information
        key = event.key()
        text = event.text()
        modifiers = event.modifiers()
        
        # Handle visual mode specially
        if self.visual_mode:
            return self._handle_visual_mode(key, text, modifiers)
            
        # Handle count prefix (numbers before commands)
        if not self.command_mode and text.isdigit() and self.count_prefix != "0":  # Don't start with 0
            self.count_prefix += text
            return True
            
        # Get count (default to 1 if no prefix)
        count = int(self.count_prefix) if self.count_prefix else 1
        
        # Handle command mode (after : key)
        if self.command_mode:
            if key == Qt.Key.Key_Escape:
                self.command_mode = False
                self.current_command = ""
                return True
            elif key == Qt.Key.Key_Return or key == Qt.Key.Key_Enter:
                self._execute_command(self.current_command)
                self.command_mode = False
                self.current_command = ""
                return True
            elif key == Qt.Key.Key_Backspace:
                self.current_command = self.current_command[:-1]
                return True
            else:
                self.current_command += text
                return True
        
        # Enter visual mode with v
        if key == Qt.Key.Key_V:
            self._enter_visual_mode()
            self.count_prefix = ""
            return True
                
        # Handle normal mode keys
        if key == Qt.Key.Key_J:  # j - move down
            self._scroll_down(count)
            self.count_prefix = ""
            return True
            
        elif key == Qt.Key.Key_K:  # k - move up
            self._scroll_up(count)
            self.count_prefix = ""
            return True
            
        elif key == Qt.Key.Key_G:  # g - go to top/bottom
            if modifiers & Qt.KeyboardModifier.ShiftModifier:  # G - go to bottom
                self._scroll_to_bottom()
            else:  # g - go to top
                self._scroll_to_top()
            self.count_prefix = ""
            return True
            
        elif key == Qt.Key.Key_D:  # d - half page down
            if modifiers & Qt.KeyboardModifier.ControlModifier:  # Ctrl+d
                self._scroll_half_page_down()
                self.count_prefix = ""
                return True
                
        elif key == Qt.Key.Key_U:  # u - half page up
            if modifiers & Qt.KeyboardModifier.ControlModifier:  # Ctrl+u
                self._scroll_half_page_up()
                self.count_prefix = ""
                return True
                
        elif key == Qt.Key.Key_F:  # f - page down
            if modifiers & Qt.KeyboardModifier.ControlModifier:  # Ctrl+f
                self._scroll_page_down()
                self.count_prefix = ""
                return True
                
        elif key == Qt.Key.Key_B:  # b - page up
            if modifiers & Qt.KeyboardModifier.ControlModifier:  # Ctrl+b
                self._scroll_page_up()
                self.count_prefix = ""
                return True
                
        elif key == Qt.Key.Key_Slash:  # / - search
            # TODO: Implement search functionality
            self.count_prefix = ""
            return True
            
        elif key == Qt.Key.Key_Colon:  # : - command mode
            self.command_mode = True
            self.current_command = ""
            self.count_prefix = ""
            return True
            
        elif key == Qt.Key.Key_Escape:  # ESC - clear state
            self.count_prefix = ""
            return True
            
        # If we've gotten this far and still have a count_prefix, it wasn't used,
        # so we should clear it
        self.count_prefix = ""
        return False
    
    def _handle_visual_mode(self, key, text, modifiers):
        """Handle key events while in visual mode."""
        # Exit visual mode with Escape
        if key == Qt.Key.Key_Escape:
            self._exit_visual_mode()
            return True
            
        # Extract text with 'e'
        if key == Qt.Key.Key_E:
            self._extract_selected_text()
            self._exit_visual_mode()
            return True
            
        # Flashcard with 'f'
        if key == Qt.Key.Key_F:
            self._create_flashcard()
            self._exit_visual_mode()
            return True
            
        # Cloze deletion with 'c'
        if key == Qt.Key.Key_C:
            self._create_cloze_deletion()
            self._exit_visual_mode()
            return True
            
        # Movement keys to extend selection
        if key == Qt.Key.Key_H:  # left
            self._extend_selection_left()
            return True
            
        if key == Qt.Key.Key_J:  # down
            self._extend_selection_down()
            return True
            
        if key == Qt.Key.Key_K:  # up
            self._extend_selection_up()
            return True
            
        if key == Qt.Key.Key_L:  # right
            self._extend_selection_right()
            return True
            
        if key == Qt.Key.Key_W:  # word forward
            self._extend_selection_word_forward()
            return True
            
        if key == Qt.Key.Key_B:  # word backward
            self._extend_selection_word_backward()
            return True
            
        return True  # Consume all keys in visual mode
    
    def _enter_visual_mode(self):
        """Enter visual mode for text selection."""
        self.visual_mode = True
        logger.debug("Entering visual mode")
        
        content_edit = self.document_view.content_edit
        
        # Setup based on content type
        if isinstance(content_edit, QTextEdit):
            # For text edit, ensure cursor is visible
            cursor = content_edit.textCursor()
            self.selection_start = cursor.position()
            content_edit.ensureCursorVisible()
            
            # Set a different cursor shape to indicate visual mode
            content_edit.viewport().setCursor(Qt.CursorShape.IBeamCursor)
            
        elif HAS_WEBENGINE and isinstance(content_edit, QWebEngineView):
            # For web view, set cursor style with JavaScript
            script = """
            document.body.style.cursor = 'text';
            
            // Create selection markers
            const marker = document.createElement('div');
            marker.id = 'vim-cursor-marker';
            marker.style.position = 'absolute';
            marker.style.width = '2px';
            marker.style.height = '1.2em';
            marker.style.backgroundColor = 'rgba(255, 0, 0, 0.7)';
            marker.style.zIndex = '9999';
            
            // Add to document
            document.body.appendChild(marker);
            
            // Position at beginning of first text node we can find
            const textNodes = [];
            function findTextNodes(node) {
                if (node.nodeType === Node.TEXT_NODE && node.textContent.trim().length > 0) {
                    textNodes.push(node);
                } else {
                    for (let i = 0; i < node.childNodes.length; i++) {
                        findTextNodes(node.childNodes[i]);
                    }
                }
            }
            
            findTextNodes(document.body);
            
            // If we found any text nodes, create a selection
            if (textNodes.length > 0) {
                const range = document.createRange();
                range.setStart(textNodes[0], 0);
                range.setEnd(textNodes[0], 0);
                
                const selection = window.getSelection();
                selection.removeAllRanges();
                selection.addRange(range);
                
                // Position marker at selection start
                const rect = range.getBoundingClientRect();
                marker.style.left = rect.left + 'px';
                marker.style.top = rect.top + 'px';
                
                // Store initial position in a way we can access later
                window.vimVisualModeStart = {
                    node: textNodes[0],
                    offset: 0
                };
            }
            """
            content_edit.page().runJavaScript(script)
        
        # Update the mode display
        self.document_view.vim_status_label.setText("Vim Mode: Visual")
        
    def _exit_visual_mode(self):
        """Exit visual mode."""
        self.visual_mode = False
        logger.debug("Exiting visual mode")
        
        content_edit = self.document_view.content_edit
        
        # Reset cursor/selection based on content type
        if isinstance(content_edit, QTextEdit):
            # Reset cursor shape
            content_edit.viewport().setCursor(Qt.CursorShape.ArrowCursor)
            
        elif HAS_WEBENGINE and isinstance(content_edit, QWebEngineView):
            # Reset cursor and remove visual markers with JavaScript
            script = """
            document.body.style.cursor = 'auto';
            
            // Remove any selection marker
            const marker = document.getElementById('vim-cursor-marker');
            if (marker) {
                marker.parentNode.removeChild(marker);
            }
            
            // Clear selection
            window.getSelection().removeAllRanges();
            delete window.vimVisualModeStart;
            """
            content_edit.page().runJavaScript(script)
        
        # Update mode display
        self.document_view.vim_status_label.setText("Vim Mode: Normal")
    
    def _scroll_down(self, count=1):
        """Scroll down by count lines."""
        content_edit = self.document_view.content_edit
        
        # Handle different widget types
        if isinstance(content_edit, QTextEdit):
            scrollbar = content_edit.verticalScrollBar()
            if scrollbar:
                # Use line height as scroll unit
                line_height = content_edit.fontMetrics().height()
                scrollbar.setValue(scrollbar.value() + count * line_height)
        elif HAS_WEBENGINE and isinstance(content_edit, QWebEngineView):
            # For web view, use JavaScript to scroll
            script = f"window.scrollBy(0, {count * 30});"  # 30px per line
            content_edit.page().runJavaScript(script)
        
    def _scroll_up(self, count=1):
        """Scroll up by count lines."""
        content_edit = self.document_view.content_edit
        
        # Handle different widget types
        if isinstance(content_edit, QTextEdit):
            scrollbar = content_edit.verticalScrollBar()
            if scrollbar:
                # Use line height as scroll unit
                line_height = content_edit.fontMetrics().height()
                scrollbar.setValue(scrollbar.value() - count * line_height)
        elif HAS_WEBENGINE and isinstance(content_edit, QWebEngineView):
            # For web view, use JavaScript to scroll
            script = f"window.scrollBy(0, -{count * 30});"  # 30px per line
            content_edit.page().runJavaScript(script)
    
    def _scroll_to_top(self):
        """Scroll to the top of the document."""
        content_edit = self.document_view.content_edit
        
        if isinstance(content_edit, QTextEdit):
            # For text edit, move cursor to start and ensure visible
            cursor = content_edit.textCursor()
            cursor.movePosition(QTextCursor.MoveOperation.Start)
            content_edit.setTextCursor(cursor)
            content_edit.ensureCursorVisible()
        elif HAS_WEBENGINE and isinstance(content_edit, QWebEngineView):
            # For web view, use JavaScript to scroll to top
            script = "window.scrollTo(0, 0);"
            content_edit.page().runJavaScript(script)
    
    def _scroll_to_bottom(self):
        """Scroll to the bottom of the document."""
        content_edit = self.document_view.content_edit
        
        if isinstance(content_edit, QTextEdit):
            # For text edit, move cursor to end and ensure visible
            cursor = content_edit.textCursor()
            cursor.movePosition(QTextCursor.MoveOperation.End)
            content_edit.setTextCursor(cursor)
            content_edit.ensureCursorVisible()
        elif HAS_WEBENGINE and isinstance(content_edit, QWebEngineView):
            # For web view, use JavaScript to scroll to bottom
            script = "window.scrollTo(0, document.body.scrollHeight);"
            content_edit.page().runJavaScript(script)
    
    def _scroll_half_page_down(self):
        """Scroll down half a page (Ctrl+d in Vim)."""
        content_edit = self.document_view.content_edit
        
        if isinstance(content_edit, QTextEdit):
            scrollbar = content_edit.verticalScrollBar()
            if scrollbar:
                # Calculate half page height
                page_step = scrollbar.pageStep()
                scrollbar.setValue(scrollbar.value() + page_step // 2)
        elif HAS_WEBENGINE and isinstance(content_edit, QWebEngineView):
            # For web view, use JavaScript to get window height and scroll
            script = "window.scrollBy(0, window.innerHeight / 2);"
            content_edit.page().runJavaScript(script)
    
    def _scroll_half_page_up(self):
        """Scroll up half a page (Ctrl+u in Vim)."""
        content_edit = self.document_view.content_edit
        
        if isinstance(content_edit, QTextEdit):
            scrollbar = content_edit.verticalScrollBar()
            if scrollbar:
                # Calculate half page height
                page_step = scrollbar.pageStep()
                scrollbar.setValue(scrollbar.value() - page_step // 2)
        elif HAS_WEBENGINE and isinstance(content_edit, QWebEngineView):
            # For web view, use JavaScript to get window height and scroll
            script = "window.scrollBy(0, -window.innerHeight / 2);"
            content_edit.page().runJavaScript(script)
    
    def _scroll_page_down(self):
        """Scroll down a full page (Ctrl+f in Vim)."""
        content_edit = self.document_view.content_edit
        
        if isinstance(content_edit, QTextEdit):
            scrollbar = content_edit.verticalScrollBar()
            if scrollbar:
                # Use page step
                page_step = scrollbar.pageStep()
                scrollbar.setValue(scrollbar.value() + page_step)
        elif HAS_WEBENGINE and isinstance(content_edit, QWebEngineView):
            # For web view, use JavaScript to get window height and scroll
            script = "window.scrollBy(0, window.innerHeight);"
            content_edit.page().runJavaScript(script)
    
    def _scroll_page_up(self):
        """Scroll up a full page (Ctrl+b in Vim)."""
        content_edit = self.document_view.content_edit
        
        if isinstance(content_edit, QTextEdit):
            scrollbar = content_edit.verticalScrollBar()
            if scrollbar:
                # Use page step
                page_step = scrollbar.pageStep()
                scrollbar.setValue(scrollbar.value() - page_step)
        elif HAS_WEBENGINE and isinstance(content_edit, QWebEngineView):
            # For web view, use JavaScript to get window height and scroll
            script = "window.scrollBy(0, -window.innerHeight);"
            content_edit.page().runJavaScript(script)
    
    def _execute_command(self, command):
        """Execute a Vim-like command."""
        command = command.strip()
        
        if not command:
            return
            
        logger.debug(f"Executing Vim command: {command}")
        
        # Handle basic commands
        if command in ['q', 'quit']:
            # Close the document
            self.document_view.close()
        elif command.startswith('set '):
            # Handle settings
            setting = command[4:].strip()
            if setting == 'novim':
                self.toggle_vim_mode()
            elif setting == 'vim':
                self.vim_mode = True
        # Add more commands as needed
    
    def _extract_selected_text(self):
        """Extract the currently selected text."""
        content_edit = self.document_view.content_edit
        
        if isinstance(content_edit, QTextEdit):
            selected_text = content_edit.textCursor().selectedText()
            if selected_text:
                self.document_view.selected_text = selected_text
                self.document_view._on_create_extract()
                
        elif HAS_WEBENGINE and isinstance(content_edit, QWebEngineView):
            # Use JavaScript to get selected text
            script = """
            window.getSelection().toString();
            """
            content_edit.page().runJavaScript(
                script,
                lambda text: self._handle_extract_from_webview(text)
            )
    
    def _handle_extract_from_webview(self, text):
        """Handle extract creation from web view selected text."""
        if text and text.strip():
            self.document_view.selected_text = text.strip()
            self.document_view._on_create_extract()
    
    def _create_flashcard(self):
        """Create a flashcard from selected text."""
        # This would require integration with a flashcard system
        # For now, we'll just extract the text
        self._extract_selected_text()
        
        # TODO: Add flashcard-specific handling when a flashcard system is available
        logger.debug("Flashcard creation requested - extracting text for now")
    
    def _create_cloze_deletion(self):
        """Create a cloze deletion from selected text."""
        # This would require integration with a flashcard/cloze system
        # For now, we'll just extract the text
        self._extract_selected_text()
        
        # TODO: Add cloze-specific handling when a cloze system is available
        logger.debug("Cloze deletion requested - extracting text for now")
    
    def _extend_selection_left(self):
        """Extend selection to the left."""
        content_edit = self.document_view.content_edit
        
        if isinstance(content_edit, QTextEdit):
            cursor = content_edit.textCursor()
            cursor.movePosition(QTextCursor.MoveOperation.Left, QTextCursor.MoveMode.KeepAnchor)
            content_edit.setTextCursor(cursor)
            
        elif HAS_WEBENGINE and isinstance(content_edit, QWebEngineView):
            script = """
            const selection = window.getSelection();
            if (selection.rangeCount > 0) {
                const range = selection.getRangeAt(0);
                
                // Try to extend one character left
                try {
                    const startNode = range.startContainer;
                    const startOffset = range.startOffset;
                    
                    if (startOffset > 0) {
                        // Can move left within current node
                        range.setStart(startNode, startOffset - 1);
                    } else {
                        // Need to find previous text node
                        // This is simplified - would need more complex traversal for complete implementation
                    }
                    
                    // Update selection
                    selection.removeAllRanges();
                    selection.addRange(range);
                    
                    // Update marker
                    const marker = document.getElementById('vim-cursor-marker');
                    if (marker) {
                        const rect = range.getBoundingClientRect();
                        marker.style.left = rect.left + 'px';
                        marker.style.top = rect.top + 'px';
                    }
                } catch (e) {
                    console.error('Error extending selection left:', e);
                }
            }
            return window.getSelection().toString();
            """
            content_edit.page().runJavaScript(script)
    
    def _extend_selection_right(self):
        """Extend selection to the right."""
        content_edit = self.document_view.content_edit
        
        if isinstance(content_edit, QTextEdit):
            cursor = content_edit.textCursor()
            cursor.movePosition(QTextCursor.MoveOperation.Right, QTextCursor.MoveMode.KeepAnchor)
            content_edit.setTextCursor(cursor)
            
        elif HAS_WEBENGINE and isinstance(content_edit, QWebEngineView):
            script = """
            const selection = window.getSelection();
            if (selection.rangeCount > 0) {
                const range = selection.getRangeAt(0);
                
                // Try to extend one character right
                try {
                    const endNode = range.endContainer;
                    const endOffset = range.endOffset;
                    
                    if (endNode.nodeType === Node.TEXT_NODE && endOffset < endNode.textContent.length) {
                        // Can move right within current text node
                        range.setEnd(endNode, endOffset + 1);
                    } else {
                        // Need to find next text node (simplified)
                    }
                    
                    // Update selection
                    selection.removeAllRanges();
                    selection.addRange(range);
                } catch (e) {
                    console.error('Error extending selection right:', e);
                }
            }
            return window.getSelection().toString();
            """
            content_edit.page().runJavaScript(script)
    
    def _extend_selection_up(self):
        """Extend selection upward."""
        content_edit = self.document_view.content_edit
        
        if isinstance(content_edit, QTextEdit):
            cursor = content_edit.textCursor()
            cursor.movePosition(QTextCursor.MoveOperation.Up, QTextCursor.MoveMode.KeepAnchor)
            content_edit.setTextCursor(cursor)
            
        elif HAS_WEBENGINE and isinstance(content_edit, QWebEngineView):
            # This is quite complex in a web page - we'll use a simplified approach
            script = """
            const selection = window.getSelection();
            if (selection.rangeCount === 0) return '';
            
            // Get current range
            const range = selection.getRangeAt(0);
            
            // Create a point slightly above the start of current selection
            const startRect = range.getBoundingClientRect();
            const x = startRect.left;
            const y = startRect.top - 20; // Move up by approximate line height
            
            // Use document.caretPositionFromPoint or document.elementFromPoint
            let newNode, newOffset;
            
            if (document.caretPositionFromPoint) {
                const pos = document.caretPositionFromPoint(x, y);
                if (pos) {
                    newNode = pos.offsetNode;
                    newOffset = pos.offset;
                }
            } else if (document.elementFromPoint) {
                // Fallback method - less accurate
                const element = document.elementFromPoint(x, y);
                if (element && element.firstChild) {
                    newNode = element.firstChild;
                    newOffset = 0;
                }
            }
            
            // If we found a node, update selection
            if (newNode) {
                try {
                    // Create new range from new point to current end
                    const newRange = document.createRange();
                    newRange.setStart(newNode, newOffset);
                    newRange.setEnd(range.endContainer, range.endOffset);
                    
                    // Apply new selection
                    selection.removeAllRanges();
                    selection.addRange(newRange);
                } catch (e) {
                    console.error('Error extending selection up:', e);
                }
            }
            
            return selection.toString();
            """
            content_edit.page().runJavaScript(script)
    
    def _extend_selection_down(self):
        """Extend selection downward."""
        content_edit = self.document_view.content_edit
        
        if isinstance(content_edit, QTextEdit):
            cursor = content_edit.textCursor()
            cursor.movePosition(QTextCursor.MoveOperation.Down, QTextCursor.MoveMode.KeepAnchor)
            content_edit.setTextCursor(cursor)
            
        elif HAS_WEBENGINE and isinstance(content_edit, QWebEngineView):
            # Similar approach to _extend_selection_up but in opposite direction
            script = """
            const selection = window.getSelection();
            if (selection.rangeCount === 0) return '';
            
            // Get current range
            const range = selection.getRangeAt(0);
            
            // Create a point slightly below the end of current selection
            const endRect = range.getBoundingClientRect();
            const x = endRect.right;
            const y = endRect.bottom + 20; // Move down by approximate line height
            
            // Use document.caretPositionFromPoint or document.elementFromPoint
            let newNode, newOffset;
            
            if (document.caretPositionFromPoint) {
                const pos = document.caretPositionFromPoint(x, y);
                if (pos) {
                    newNode = pos.offsetNode;
                    newOffset = pos.offset;
                }
            } else if (document.elementFromPoint) {
                // Fallback method - less accurate
                const element = document.elementFromPoint(x, y);
                if (element && element.firstChild) {
                    newNode = element.firstChild;
                    newOffset = 0;
                }
            }
            
            // If we found a node, update selection
            if (newNode) {
                try {
                    // Create new range from current start to new point
                    const newRange = document.createRange();
                    newRange.setStart(range.startContainer, range.startOffset);
                    newRange.setEnd(newNode, newOffset);
                    
                    // Apply new selection
                    selection.removeAllRanges();
                    selection.addRange(newRange);
                } catch (e) {
                    console.error('Error extending selection down:', e);
                }
            }
            
            return selection.toString();
            """
            content_edit.page().runJavaScript(script)
    
    def _extend_selection_word_forward(self):
        """Extend selection to the next word."""
        content_edit = self.document_view.content_edit
        
        if isinstance(content_edit, QTextEdit):
            cursor = content_edit.textCursor()
            cursor.movePosition(QTextCursor.MoveOperation.NextWord, QTextCursor.MoveMode.KeepAnchor)
            content_edit.setTextCursor(cursor)
            
        elif HAS_WEBENGINE and isinstance(content_edit, QWebEngineView):
            # Simplified word selection - better implementations would use text node traversal
            script = """
            const selection = window.getSelection();
            if (!selection.rangeCount) return '';
            
            const range = selection.getRangeAt(0);
            const endNode = range.endContainer;
            
            if (endNode.nodeType === Node.TEXT_NODE) {
                const text = endNode.textContent;
                const endOffset = range.endOffset;
                
                // Find next word boundary
                let nextSpace = text.indexOf(' ', endOffset);
                if (nextSpace === -1) nextSpace = text.length;
                
                // Extend to word boundary
                range.setEnd(endNode, nextSpace + 1);
                
                // Update selection
                selection.removeAllRanges();
                selection.addRange(range);
            }
            
            return selection.toString();
            """
            content_edit.page().runJavaScript(script)
    
    def _extend_selection_word_backward(self):
        """Extend selection to the previous word."""
        content_edit = self.document_view.content_edit
        
        if isinstance(content_edit, QTextEdit):
            cursor = content_edit.textCursor()
            cursor.movePosition(QTextCursor.MoveOperation.PreviousWord, QTextCursor.MoveMode.KeepAnchor)
            content_edit.setTextCursor(cursor)
            
        elif HAS_WEBENGINE and isinstance(content_edit, QWebEngineView):
            # Simplified word selection - better implementations would use text node traversal
            script = """
            const selection = window.getSelection();
            if (!selection.rangeCount) return '';
            
            const range = selection.getRangeAt(0);
            const startNode = range.startContainer;
            
            if (startNode.nodeType === Node.TEXT_NODE) {
                const text = startNode.textContent;
                const startOffset = range.startOffset;
                
                // Find previous word boundary
                const textBefore = text.substring(0, startOffset);
                let prevSpace = textBefore.lastIndexOf(' ');
                
                // If no space found, go to beginning of text
                if (prevSpace === -1) prevSpace = 0;
                else prevSpace += 1; // Move past the space
                
                // Extend selection
                range.setStart(startNode, prevSpace);
                
                // Update selection
                selection.removeAllRanges();
                selection.addRange(range);
            }
            
            return selection.toString();
            """
            content_edit.page().runJavaScript(script)

class DocumentView(QWidget):
    """UI component for viewing and processing documents."""
    
    extractCreated = pyqtSignal(int)  # extract_id
    
    def __init__(self, db_session, document_id=None):
        super().__init__()
        
        self.db_session = db_session
        self.document_id = document_id
        self.document = None
        self.selected_text = ""
        self.content_text = ""
        self.youtube_callback = None
        
        # View state tracking
        self.view_state = {
            "zoom_factor": 1.0,
            "position": None,
            "size": None,
            "scroll_position": 0
        }
        
        # Create the Vim key handler
        self.vim_handler = VimKeyHandler(self)
        
        # Create the UI
        self._create_ui()
        
        # Load document if provided
        if document_id:
            self.load_document(document_id)
    
    def _create_ui(self):
        """Create the UI components."""
        # Main layout
        layout = QVBoxLayout()
        layout.setContentsMargins(0, 0, 0, 0)
        
        # Toolbar
        toolbar = QToolBar()
        
        # Document navigation actions
        self.prev_action = QAction("Previous", self)
        self.prev_action.triggered.connect(self._on_previous)
        toolbar.addAction(self.prev_action)
        
        self.next_action = QAction("Next", self)
        self.next_action.triggered.connect(self._on_next)
        toolbar.addAction(self.next_action)
        
        toolbar.addSeparator()
        
        # Extract actions
        self.create_extract_action = QAction("Create Extract", self)
        self.create_extract_action.triggered.connect(self._on_create_extract)
        toolbar.addAction(self.create_extract_action)
        
        toolbar.addSeparator()
        
        # Vim mode toggle
        self.vim_toggle_action = QAction("Vim Mode", self)
        self.vim_toggle_action.setCheckable(True)
        self.vim_toggle_action.setChecked(self.vim_handler.vim_mode)
        self.vim_toggle_action.triggered.connect(self._toggle_vim_mode)
        toolbar.addAction(self.vim_toggle_action)
        
        layout.addWidget(toolbar)
        
        # Content splitter
        splitter = QSplitter(Qt.Orientation.Horizontal)
        
        # Document content area
        self.content_widget = QWidget()
        self.content_layout = QVBoxLayout(self.content_widget)
        self.content_layout.setContentsMargins(0, 0, 0, 0)
        self.content_layout.setSpacing(0)  # Minimize spacing between elements
        
        # Content will be added dynamically based on document type
        
        splitter.addWidget(self.content_widget)
        
        # Extracts area
        self.extract_view = DocumentExtractsView(self.db_session)
        self.extract_view.extractSelected.connect(self._on_extract_selected)
        splitter.addWidget(self.extract_view)
        
        # Set initial sizes - give almost all space to the PDF content (95/5 split)
        splitter.setSizes([950, 50])
        
        # Make sure the split ratio is maintained when resizing
        splitter.setStretchFactor(0, 19)  # Content gets 19x the stretch of extracts
        splitter.setStretchFactor(1, 1)   # Extracts get 1x stretch factor
        
        layout.addWidget(splitter, 1)  # Add stretch factor of 1 to expand
        
        # Add Vim mode status bar - use minimal height
        self.vim_status_widget = QWidget()
        self.vim_status_widget.setMaximumHeight(20)  # Limit height
        vim_status_layout = QHBoxLayout(self.vim_status_widget)
        vim_status_layout.setContentsMargins(2, 1, 2, 1)  # Minimal margins
        
        self.vim_status_label = QLabel("Vim Mode: Normal")
        vim_status_layout.addWidget(self.vim_status_label)
        
        # Add vim command display
        self.vim_command_label = QLabel("")
        vim_status_layout.addWidget(self.vim_command_label)
        
        # Add visual mode indicator
        self.vim_visual_label = QLabel("")
        self.vim_visual_label.setStyleSheet("color: #c22;")  # Red text for visual mode
        vim_status_layout.addWidget(self.vim_visual_label)
        
        vim_status_layout.addStretch(1)  # Push everything to the left
        
        # Add keyboard shortcuts info
        self.vim_shortcuts_label = QLabel("v:visual e:extract f:flashcard c:cloze")
        self.vim_shortcuts_label.setStyleSheet("color: #666; font-size: 9pt;")
        vim_status_layout.addWidget(self.vim_shortcuts_label)
        
        # Add to main layout - no stretch (keep minimal)
        layout.addWidget(self.vim_status_widget, 0)  # Use 0 stretch factor
        
        self.setLayout(layout)
        
        # Update Vim status bar visibility based on mode
        self._update_vim_status_visibility()
    
    def _update_vim_status_visibility(self):
        """Update the visibility of the Vim status bar based on mode."""
        self.vim_status_widget.setVisible(self.vim_handler.vim_mode)
        self.vim_toggle_action.setChecked(self.vim_handler.vim_mode)
    
    def _toggle_vim_mode(self):
        """Toggle Vim mode on/off."""
        is_enabled = self.vim_handler.toggle_vim_mode()
        self.vim_toggle_action.setChecked(is_enabled)
        self._update_vim_status_visibility()
    
    def keyPressEvent(self, event):
        """Handle key press events for Vim-like navigation."""
        # Check if Vim handler wants to handle this key
        if self.vim_handler.handle_key_event(event):
            # Update command display if in command mode
            if self.vim_handler.command_mode:
                self.vim_status_label.setText("Vim Mode: Command")
                self.vim_command_label.setText(":" + self.vim_handler.current_command)
                self.vim_visual_label.setText("")
            elif self.vim_handler.visual_mode:
                self.vim_status_label.setText("Vim Mode: Visual")
                self.vim_command_label.setText("")
                self.vim_visual_label.setText("[Selection Mode]")
            else:
                self.vim_status_label.setText("Vim Mode: Normal")
                self.vim_visual_label.setText("")
                # Show count prefix if any
                if self.vim_handler.count_prefix:
                    self.vim_command_label.setText(self.vim_handler.count_prefix)
                else:
                    self.vim_command_label.setText("")
            
            # Event was handled
            return
        
        # If not handled by Vim mode, pass to parent
        super().keyPressEvent(event)
    
    def _create_webview_and_setup(self, html_content, base_url):
        """Create a QWebEngineView and set it up with the content."""
        if not HAS_WEBENGINE:
            logger.warning("QWebEngineView not available, falling back to QTextEdit")
            editor = QTextEdit()
            editor.setReadOnly(True)
            editor.setHtml(html_content)
            editor.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
            return editor
        
        # Create WebEngine view for HTML content
        webview = QWebEngineView()
        
        # Set size policy to expand
        webview.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        
        # Configure settings to allow all JavaScript functionality
        settings = webview.settings()
        settings.setAttribute(settings.WebAttribute.JavascriptEnabled, True)
        settings.setAttribute(settings.WebAttribute.JavascriptCanAccessClipboard, True)
        settings.setAttribute(settings.WebAttribute.JavascriptCanOpenWindows, True)
        settings.setAttribute(settings.WebAttribute.LocalContentCanAccessRemoteUrls, True)
        settings.setAttribute(settings.WebAttribute.AllowRunningInsecureContent, True)
        settings.setAttribute(settings.WebAttribute.PluginsEnabled, True)
        
        # Create a channel for JavaScript communication
        channel = QWebChannel(webview.page())
        
        # Create callback handler
        callback_handler = WebViewCallback(self)
        channel.registerObject("callbackHandler", callback_handler)
        
        # Set the channel to the page
        webview.page().setWebChannel(channel)
        
        # Check if we need to inject JavaScript libraries
        injected_html = self._inject_javascript_libraries(html_content)
        
        # Load content with proper base URL for resources
        if base_url:
            webview.setHtml(injected_html, base_url)
        else:
            webview.setHtml(injected_html)
        
        # Inject JavaScript to capture selections - using a cleaner approach
        selection_js = """
        // Wait for document to be fully loaded
        document.addEventListener('DOMContentLoaded', function() {
            // Add selection change listener
            document.addEventListener('selectionchange', function() {
                const selection = window.getSelection();
                const text = selection.toString();
                // Only send non-empty selections
                if (text && text.trim().length > 0) {
                    // Check if callback handler is available
                    if (typeof window.callbackHandler !== 'undefined') {
                        window.callbackHandler.selectionChanged(text);
                    } else {
                        console.error('Callback handler not available');
                    }
                }
            });

            // Track scroll position changes
            document.addEventListener('scroll', function() {
                // Throttle scroll events to reduce overhead
                if (window.scrollTrackTimeout) {
                    clearTimeout(window.scrollTrackTimeout);
                }
                window.scrollTrackTimeout = setTimeout(function() {
                    const scrollPosition = window.pageYOffset || document.documentElement.scrollTop;
                    if (typeof window.callbackHandler !== 'undefined') {
                        // Store position in a window variable so we can access it later
                        window.lastScrollPosition = scrollPosition;
                    }
                }, 200); // 200ms throttle
            });
            
            // Initialize any custom libraries
            if (typeof initializeCustomLibraries === 'function') {
                try {
                    initializeCustomLibraries();
                    console.log('Custom libraries initialized');
                } catch (e) {
                    console.error('Error initializing custom libraries:', e);
                }
            }
        });
        """
        
        # Simple direct selection capture script as fallback
        simple_selection_js = """
        document.onselectionchange = function() {
            var selection = window.getSelection();
            var text = selection.toString();
            if (text && text.trim().length > 0) {
                window.text_selection = text; // Store in a global variable
            }
        };
        """
        
        # Inject the simple selection script immediately
        webview.page().runJavaScript(simple_selection_js)
        
        # Inject the main script after the page has loaded
        webview.loadFinished.connect(lambda ok: webview.page().runJavaScript(selection_js))
        
        # Add a method to manually get the current selection
        def check_selection():
            webview.page().runJavaScript(
                "window.getSelection().toString() || window.text_selection || '';",
                self._handle_webview_selection
            )
        
        # Store the method to check selection
        webview.check_selection = check_selection
        
        # Connect mouse release to check selection
        webview.mouseReleaseEvent = lambda event: (
            super(QWebEngineView, webview).mouseReleaseEvent(event),
            check_selection()
        )
        
        # Connect context menu request to check selection before showing menu
        webview.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        
        # Store the original handler
        original_handler = self._on_content_menu
        
        # Create a wrapper that first checks selection
        def context_menu_wrapper(pos):
            check_selection()
            # Call the original handler after a small delay
            QApplication.processEvents()
            original_handler(pos)
        
        webview.customContextMenuRequested.connect(context_menu_wrapper)
        
        return webview
    
    def _inject_javascript_libraries(self, html_content):
        """Inject common JavaScript libraries into HTML content based on content needs."""
        
        # Check if the document already has a head tag
        has_head = "<head>" in html_content
        has_body = "<body>" in html_content
        
        # Library CDN URLs
        libraries = {
            'markdown': "https://cdn.jsdelivr.net/npm/marked/marked.min.js",
            'mermaid': "https://cdn.jsdelivr.net/npm/mermaid/dist/mermaid.min.js",
            'plotly': "https://cdn.plot.ly/plotly-latest.min.js",
            'katex': "https://cdn.jsdelivr.net/npm/katex@0.16.8/dist/katex.min.js",
            'katex-css': "https://cdn.jsdelivr.net/npm/katex@0.16.8/dist/katex.min.css",
            'katex-autorender': "https://cdn.jsdelivr.net/npm/katex@0.16.8/dist/contrib/auto-render.min.js",
            'three': "https://cdn.jsdelivr.net/npm/three@0.157.0/build/three.min.js"
        }
        
        # Determine which libraries to inject based on content
        to_inject = []
        
        # Check for library markers in content
        if "```mermaid" in html_content or "mermaid" in html_content:
            to_inject.append('mermaid')
        
        if "```math" in html_content or "\\(" in html_content or "\\[" in html_content or "$" in html_content or "math" in html_content.lower():
            to_inject.append('katex')
            to_inject.append('katex-css')
            to_inject.append('katex-autorender')
        
        if "```plotly" in html_content or "Plotly" in html_content:
            to_inject.append('plotly')
        
        if "```markdown" in html_content or "marked" in html_content or "markdown" in html_content.lower():
            to_inject.append('markdown')
        
        if "three.js" in html_content or "THREE." in html_content:
            to_inject.append('three')
        
        # Return original content if no libraries needed or if we can't find insertion points
        if not to_inject or (not has_head and not has_body):
            return html_content
        
        # Build script tags
        scripts = ""
        styles = ""
        
        for lib in to_inject:
            if lib.endswith('-css'):
                styles += f'<link rel="stylesheet" href="{libraries[lib]}">\n'
            else:
                scripts += f'<script src="{libraries[lib]}"></script>\n'
        
        # Add initialization script
        init_script = """
<script>
function initializeCustomLibraries() {
    // Initialize mermaid if available
    if (typeof mermaid !== 'undefined') {
        try {
            mermaid.initialize({
                startOnLoad: true,
                theme: 'neutral'
            });
        } catch (e) {
            console.error('Mermaid init error:', e);
        }
    }
    
    // Initialize KaTeX if available
    if (typeof katex !== 'undefined') {
        console.log('KaTeX detected, initializing...');
        
        // First approach: Find specific elements with class names
        document.querySelectorAll('.math, .katex').forEach(element => {
            try {
                // Get the raw text and trim whitespace
                let texContent = element.textContent.trim();
                console.log('Processing KaTeX element:', texContent.substring(0, 30) + '...');
                
                // Create a new element to render into (to avoid modifying the original)
                let renderElement = document.createElement('div');
                renderElement.className = 'katex-output';
                element.innerHTML = ''; // Clear the original content
                element.appendChild(renderElement);
                
                // Render the KaTeX
                katex.render(texContent, renderElement, {
                    throwOnError: false,
                    displayMode: element.classList.contains('display')
                });
            } catch (e) {
                console.error('KaTeX rendering error:', e);
                element.innerHTML = '<span style="color:red">KaTeX Error: ' + e.message + '</span><br>' + element.textContent;
            }
        });
        
        // Render inline math
        document.querySelectorAll('.math-inline').forEach(element => {
            try {
                let texContent = element.textContent.trim();
                katex.render(texContent, element, {
                    throwOnError: false,
                    displayMode: false
                });
            } catch (e) {
                console.error('KaTeX inline error:', e);
                element.innerHTML = '<span style="color:red">Error</span>';
            }
        });
        
        // Second approach: Use the auto-render extension if available
        if (typeof renderMathInElement !== 'undefined') {
            console.log('Using KaTeX auto-render extension');
            try {
                // Automatically render math in the entire document
                renderMathInElement(document.body, {
                    delimiters: [
                        {left: "$$", right: "$$", display: true},
                        {left: "$", right: "$", display: false},
                        {left: "\\(", right: "\\)", display: false},
                        {left: "\\[", right: "\\]", display: true}
                    ],
                    throwOnError: false
                });
            } catch (e) {
                console.error('KaTeX auto-render error:', e);
            }
        }
    }
    
    // Initialize markdown if available
    if (typeof marked !== 'undefined') {
        console.log('Marked.js detected, initializing...');
        
        // Set up marked options for better rendering
        marked.setOptions({
            gfm: true,          // Enable GitHub flavored markdown
            breaks: true,       // Convert line breaks to <br>
            smartLists: true,   // Use smart list behavior
            smartypants: true,  // Use smart punctuation
            xhtml: true         // Return XHTML compliant output
        });
        
        document.querySelectorAll('.markdown').forEach(element => {
            try {
                // Get the raw text
                let mdContent = element.textContent.trim();
                console.log('Processing Markdown element:', mdContent.substring(0, 30) + '...');
                
                // Parse markdown
                let htmlContent = marked.parse(mdContent);
                
                // Create a wrapper to maintain any styling
                let wrapper = document.createElement('div');
                wrapper.className = 'markdown-output';
                wrapper.innerHTML = htmlContent;
                
                // Replace content
                element.innerHTML = '';
                element.appendChild(wrapper);
            } catch (e) {
                console.error('Markdown error:', e);
                element.innerHTML = '<div class="error">Markdown rendering error: ' + e.message + '</div>';
            }
        });
    }
    
    console.log('Custom libraries initialization complete');
}

// Execute after DOM is fully loaded
document.addEventListener('DOMContentLoaded', function() {
    console.log('DOM loaded, initializing custom libraries');
    setTimeout(initializeCustomLibraries, 100);
});

// Add a button to manually trigger initialization if automatic doesn't work
window.addEventListener('load', function() {
    // Check if any library-specific elements exist but haven't been processed
    let needsInit = (
        (document.querySelector('.markdown') && !document.querySelector('.markdown-output')) ||
        (document.querySelector('.math, .katex') && !document.querySelector('.katex-output'))
    );
    
    if (needsInit) {
        console.log('Library elements found that need initialization');
        
        // Create a button to manually trigger initialization
        let btn = document.createElement('button');
        btn.textContent = 'Initialize Custom Content';
        btn.style.position = 'fixed';
        btn.style.bottom = '10px';
        btn.style.right = '10px';
        btn.style.zIndex = '9999';
        btn.style.padding = '8px 16px';
        btn.style.backgroundColor = '#0066cc';
        btn.style.color = 'white';
        btn.style.border = 'none';
        btn.style.borderRadius = '4px';
        btn.style.cursor = 'pointer';
        
        btn.onclick = function() {
            initializeCustomLibraries();
            btn.textContent = 'Initialization Complete';
            setTimeout(function() { 
                btn.style.display = 'none'; 
            }, 2000);
        };
        
        document.body.appendChild(btn);
        
        // Try one more time automatically
        setTimeout(initializeCustomLibraries, 1000);
    }
});
</script>
"""
        
        # Inject into HTML
        if has_head:
            # Insert before the closing head tag
            html_content = html_content.replace("</head>", styles + scripts + init_script + "</head>")
        elif has_body:
            # Insert after the opening body tag
            html_content = html_content.replace("<body>", "<body>\n" + styles + scripts + init_script)
        else:
            # Wrap the entire content
            html_content = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1">
    <title>Document</title>
    {styles}
    {scripts}
    {init_script}
</head>
<body>
    {html_content}
</body>
</html>
"""
        
        return html_content
    
    def keep_alive(self, obj):
        """Keep a reference to an object to prevent garbage collection."""
        if not hasattr(self, '_kept_references'):
            self._kept_references = []
        self._kept_references.append(obj)
        
    def load_document(self, document_id):
        """Load a document for viewing.
        
        Args:
            document_id (int): ID of the document to load.
        """
        try:
            # Keep any existing references alive
            if hasattr(self, 'webview'):
                self.keep_alive(self.webview)
            if hasattr(self, 'youtube_callback'):
                self.keep_alive(self.youtube_callback)
                
            # Clear the document area
            self._clear_content_layout()
            
            # Get the document
            self.document_id = document_id
            self.document = self.db_session.query(Document).filter_by(id=document_id).first()
            
            if not self.document:
                raise ValueError(f"Document not found: {document_id}")
            
            # Update extracts view
            if hasattr(self, 'extract_view') and self.extract_view:
                self.extract_view.load_extracts_for_document(document_id)
            
            # Load the document content based on its type
            doc_type = self.document.content_type.lower() if hasattr(self.document, 'content_type') and self.document.content_type else "text"
            
            logger.debug(f"Loading document: {self.document.title} (Type: {doc_type})")
            
            # Handle different document types
            if doc_type == "youtube":
                self._load_youtube()
            elif doc_type == "epub":
                self._load_epub(self.db_session, self.document)
            elif doc_type == "pdf":
                self._load_pdf()
            elif doc_type == "html" or doc_type == "htm":
                self._load_html()
            elif doc_type == "txt":
                self._load_text()
            else:
                # Default to text view
                self._load_text()
            
            # Set window title to document title
            if hasattr(self, 'setWindowTitle') and callable(self.setWindowTitle):
                self.setWindowTitle(self.document.title)
                
            return True
        except Exception as e:
            logger.exception(f"Error loading document {document_id}: {e}")
            
            # Show error in view
            error_widget = QLabel(f"Error loading document: {str(e)}")
            error_widget.setWordWrap(True)
            error_widget.setStyleSheet("color: red; padding: 20px;")
            
            self._clear_content_layout()
            if hasattr(self, 'content_layout'):
                self.content_layout.addWidget(error_widget)
                
            return False
    
    def _load_pdf(self):
        """Load and display a PDF document."""
        try:
            file_path = self.document.file_path
            
            # Check if file exists
            if not os.path.isfile(file_path):
                logger.error(f"PDF file not found: {file_path}")
                
                # Try alternative file path if in temporary directory
                if '/tmp/' in file_path:
                    tmp_dir = os.path.dirname(file_path)
                    if os.path.exists(tmp_dir):
                        files = os.listdir(tmp_dir)
                        pdf_files = [f for f in files if f.endswith('.pdf')]
                        if pdf_files:
                            new_path = os.path.join(tmp_dir, pdf_files[0])
                            logger.info(f"Using alternative PDF file found in the same directory: {new_path}")
                            file_path = new_path
                            
                            # Update the document's file_path
                            self.document.file_path = file_path
                            self.db_session.commit()
                        else:
                            logger.error(f"No PDF files found in {tmp_dir}")
                            raise FileNotFoundError(f"PDF file not found: {file_path}")
                    else:
                        logger.error(f"Temporary directory not found: {tmp_dir}")
                        raise FileNotFoundError(f"PDF file not found: {file_path}")
                else:
                    raise FileNotFoundError(f"PDF file not found: {file_path}")
            
            # First try using PyMuPDF for advanced features
            try:
                import fitz  # PyMuPDF
                
                # Custom PDF viewer using PyMuPDF
                from ui.pdf_view import PDFViewWidget
                
                # Create the PDF view widget
                pdf_widget = PDFViewWidget(self.document, self.db_session)
                
                # Connect extract created signal if available
                if hasattr(self, 'extractCreated') and hasattr(pdf_widget, 'extractCreated'):
                    pdf_widget.extractCreated.connect(self.extractCreated.emit)
                
                # Add to layout
                self.content_layout.addWidget(pdf_widget)
                
                # Store content edit for later use
                self.content_edit = pdf_widget
                
                logger.info(f"Loaded PDF with PyMuPDF: {file_path}")
                return
                
            except (ImportError, Exception) as e:
                logger.warning(f"Could not use PyMuPDF for PDF: {str(e)}. Falling back to QPdfView.")
                # Fall back to QPdfView
                pass
                
            # Fallback: Use QPdfView from Qt
            from PyQt6.QtPdf import QPdfDocument
            from PyQt6.QtPdfWidgets import QPdfView
            
            # Create PDF view
            pdf_view = QPdfView()
            
            # Create PDF document
            pdf_document = QPdfDocument()
            
            # Load the PDF file
            pdf_document.load(file_path)
            
            # Set the document to the view
            pdf_view.setDocument(pdf_document)
            
            # Add to layout
            self.content_layout.addWidget(pdf_view)
            
            # Store content edit for later use (position tracking)
            self.content_edit = pdf_view
            
            # Extract text content for context
            # This would require a PDF text extraction library
            self.content_text = "PDF content"
            
            # Restore reading position if available
            self._restore_position()
            
            logger.info(f"Loaded PDF with QPdfView: {file_path}")
            
        except ImportError:
            logger.error("PDF viewing requires PyQt6.QtPdf and PyQt6.QtPdfWidgets")
            label = QLabel("PDF viewing requires additional modules that are not installed.")
            self.content_layout.addWidget(label)
        except Exception as e:
            logger.exception(f"Error loading PDF: {e}")
            label = QLabel(f"Error loading PDF: {str(e)}")
            self.content_layout.addWidget(label)

    def _load_epub(self, db_session, document):
        """
        Load document in EPUB format.
        
        Args:
            db_session: SQLAlchemy session
            document: Document object
            
        Returns:
            bool: True if document was loaded successfully
        """
        try:
            if not HAS_WEBENGINE:
                raise Exception("Web engine not available.")
            
            # Create the web view
            content_edit = QWebEngineView()
            
            # Load the EPUB into the handler
            try:
                epub_handler = EPUBHandler()
                # Load the file data and prepare it for the webview
                content_results = epub_handler.extract_content(document.file_path)
                html_content = content_results['html']
                
                # Check if we can detect any JS libraries in the HTML content
                # For better display with certain EPUB formats
                has_jquery = "jquery" in html_content.lower()
                
                # Apply jQuery from resources if not already in the HTML
                if not has_jquery:
                    logger.debug("EPUB does not contain jQuery, will inject if needed")
                    # Handled via injectScriptTag in later methods
                
                # Set the EPUB content to the view
                content_edit.setHtml(html_content, QUrl.fromLocalFile(document.file_path))
                
                # Set up position tracking
                from ui.load_epub_helper import setup_epub_webview
                from core.utils.theme_manager import ThemeManager
                
                # Get theme manager from MainWindow or create a new instance
                app = QApplication.instance()
                main_window = None
                for widget in app.topLevelWidgets():
                    if widget.__class__.__name__ == "MainWindow":
                        main_window = widget
                        break
                
                theme_manager = None
                if main_window and hasattr(main_window, 'theme_manager'):
                    theme_manager = main_window.theme_manager
                else:
                    # Create a new instance if not available
                    from core.settings.settings_manager import SettingsManager
                    settings_manager = SettingsManager()
                    theme_manager = ThemeManager(settings_manager)
                
                # Add SuperMemo SM-18 style incremental reading capabilities
                self._add_supermemo_features(content_edit)
                
                # Set up the EPUB view with theme
                setup_epub_webview(document, content_edit, db_session, restore_position=True, theme_manager=theme_manager)
                
                # Add it to our layout
                self.content_layout.addWidget(content_edit)
                
                # Store content for later use
                self.content_edit = content_edit
                self.content_text = content_results['text']
                
                return True
                
            except Exception as e:
                logger.exception(f"Failed to load EPUB document: {e}")
                content_edit.setHtml(f"<h1>Error loading EPUB</h1><p>{str(e)}</p>")
                self.content_layout.addWidget(content_edit)
                return False
                
        except Exception as e:
            logger.exception(f"Error setting up EPUB view: {e}")
            self._set_error_message(f"Could not load EPUB document: {e}")
            return False
    
    def _load_text(self):
        """Load and display a text document."""
        try:
            # Create text edit
            text_edit = QTextEdit()
            text_edit.setReadOnly(True)
            
            # Set size policy to expand
            text_edit.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
            
            # Read file content
            with open(self.document.file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Set content
            text_edit.setPlainText(content)
            
            # Add to layout - use stretch factor to expand
            self.content_layout.addWidget(text_edit, 1)
            
            # Store for later use
            self.content_edit = text_edit
            self.content_text = content
            
            # Restore position
            self._restore_position()
            
        except Exception as e:
            logger.exception(f"Error loading text document: {e}")
            label = QLabel(f"Error loading text document: {str(e)}")
            self.content_layout.addWidget(label)
    
    def _load_html(self):
        """Load and display an HTML document."""
        try:
            # Log file info
            logger.info(f"Loading HTML document from {self.document.file_path}")
            file_size = os.path.getsize(self.document.file_path) if os.path.exists(self.document.file_path) else 0
            logger.info(f"HTML file size: {file_size} bytes")
            
            # Check if file exists and is not empty
            if not os.path.exists(self.document.file_path):
                logger.error(f"HTML file does not exist: {self.document.file_path}")
                label = QLabel(f"HTML file not found: {os.path.basename(self.document.file_path)}")
                self.content_layout.addWidget(label)
                return
                
            if file_size == 0:
                logger.error(f"HTML file is empty: {self.document.file_path}")
                label = QLabel("The HTML file is empty. No content to display.")
                self.content_layout.addWidget(label)
                return
            
            # Read HTML file
            try:
                with open(self.document.file_path, 'r', encoding='utf-8') as f:
                    html_content = f.read()
            except UnicodeDecodeError:
                # Try with a different encoding if UTF-8 fails
                logger.warning(f"UTF-8 decoding failed, trying with ISO-8859-1")
                with open(self.document.file_path, 'r', encoding='ISO-8859-1') as f:
                    html_content = f.read()
                    
            # Log content preview for debugging
            content_preview = html_content[:200] + "..." if len(html_content) > 200 else html_content
            logger.debug(f"HTML content preview: {content_preview}")
            
            if not html_content.strip():
                logger.error("HTML content is empty after reading")
                label = QLabel("The HTML file contains no content.")
                self.content_layout.addWidget(label)
                return
            
            # Parse with BeautifulSoup to extract text
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Check if there's any meaningful content
            body_content = soup.body.get_text().strip() if soup.body else ""
            self.content_text = soup.get_text()
            
            if not body_content:
                logger.warning("HTML document contains no visible text content")
                
                # Try to display the raw HTML instead
                text_edit = QTextEdit()
                text_edit.setReadOnly(True)
                text_edit.setPlainText(html_content)
                text_edit.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
                self.content_layout.addWidget(text_edit, 1)  # Add stretch factor
                self.content_edit = text_edit
                logger.info("Displaying raw HTML content instead")
                return
            
            # Set up base URL for loading resources
            base_url = QUrl.fromLocalFile(os.path.dirname(self.document.file_path) + os.path.sep)
            
            # Log which JavaScript libraries might be needed
            libs_to_check = ['mermaid', 'katex', 'plotly', 'markdown', 'three.js']
            detected_libs = []
            
            for lib in libs_to_check:
                if lib in html_content.lower():
                    detected_libs.append(lib)
                    
            if detected_libs:
                logger.info(f"Detected potential JavaScript libraries in HTML: {', '.join(detected_libs)}")
            
            # Create web view with libraries injected as needed
            webview = self._create_webview_and_setup(html_content, base_url)
            
            # Set size policy to expand
            if HAS_WEBENGINE and isinstance(webview, QWebEngineView):
                webview.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
            
            # Add to layout with stretch
            self.content_layout.addWidget(webview, 1)
            
            # Store for later use
            self.content_edit = webview
            self.web_view = webview  # Keep reference for _save_position
            
            # Set up position tracking similar to EPUB
            setup_epub_webview(self.document, webview, self.db_session)
            
            logger.info("HTML document loaded successfully")
            
        except Exception as e:
            logger.exception(f"Error loading HTML document: {e}")
            label = QLabel(f"Error loading HTML document: {str(e)}")
            self.content_layout.addWidget(label)
    
    def _load_docx(self):
        """Load and display a DOCX document."""
        try:
            # Use python-docx to extract text
            import docx
            
            doc = docx.Document(self.document.file_path)
            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Create text edit
            text_edit = QTextEdit()
            text_edit.setReadOnly(True)
            text_edit.setPlainText(content)
            
            # Add to layout
            self.content_layout.addWidget(text_edit)
            
            # Store for later use
            self.content_edit = text_edit
            self.content_text = content
            
            # Restore position
            self._restore_position()
            
        except ImportError:
            logger.error("DOCX viewing requires python-docx library")
            label = QLabel("DOCX viewing requires additional libraries that are not installed.")
            self.content_layout.addWidget(label)
        except Exception as e:
            logger.exception(f"Error loading DOCX document: {e}")
            label = QLabel(f"Error loading DOCX document: {str(e)}")
            self.content_layout.addWidget(label)

    def _load_youtube(self):
        """Load YouTube video content."""
        try:
            # Import required widgets here to ensure they're in scope
            from PyQt6.QtWidgets import QWidget, QVBoxLayout, QHBoxLayout, QPushButton, QLabel, QLineEdit, QSplitter, QSizePolicy
            from PyQt6.QtWebEngineWidgets import QWebEngineView
            from PyQt6.QtCore import Qt
            
            if not HAS_WEBENGINE:
                raise ImportError("YouTube videos require QWebEngineView which is not available")
                
            # Clear any existing content from the layout
            self._clear_content_layout()
            
            # Get video ID
            video_id = extract_video_id_from_document(self.document)
            
            if not video_id:
                logger.error("Could not extract YouTube video ID")
                raise ValueError("Could not extract YouTube video ID from document")
            
            # Create a splitter to hold both the video and transcript
            self.content_splitter = QSplitter(Qt.Orientation.Vertical)
            self.content_splitter.setChildrenCollapsible(False)  # Prevent sections from being fully collapsed
            
            # Create a container for the video
            video_container = QWidget()
            video_layout = QVBoxLayout(video_container)
            video_layout.setContentsMargins(0, 0, 0, 0)
            video_layout.setSpacing(5)
            
            # Create a WebView and make it expand to fill available space
            self.webview = QWebEngineView()
            self.webview.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)  # Allow expansion
            video_layout.addWidget(self.webview, stretch=1)  # Add stretch to prioritize video
            
            # Get the target position (if available)
            target_position = 0
            if hasattr(self.document, 'position') and self.document.position:
                target_position = self.document.position
                
            # Add controls widget with more compact layout
            controls_widget = QWidget()
            timestamp_layout = QHBoxLayout(controls_widget)
            timestamp_layout.setContentsMargins(5, 2, 5, 2)
            timestamp_layout.setSpacing(2)
            
            # Add position label
            self.position_label = QLabel(f"Position: {target_position}s")
            timestamp_layout.addWidget(self.position_label)
            
            # Add spacer between position and seek controls
            timestamp_layout.addStretch(1)
            
            # Add seek label
            timestamp_label = QLabel("Seek to:")
            timestamp_layout.addWidget(timestamp_label)
            
            # Add input field for timestamp
            self.seek_time_input = QLineEdit()
            self.seek_time_input.setPlaceholderText("Enter time in seconds or MM:SS")
            self.seek_time_input.setText(str(target_position))
            timestamp_layout.addWidget(self.seek_time_input)
            
            # Add seek button
            seek_button = QPushButton("Seek")
            seek_button.clicked.connect(self._on_seek_youtube_position)
            timestamp_layout.addWidget(seek_button)
            
            # Add a save button for manual saving of position
            save_button = QPushButton("Save Position")
            save_button.clicked.connect(self._on_save_youtube_position)
            timestamp_layout.addWidget(save_button)
            
            # Add controls to video container (no stretch, keep it compact)
            video_layout.addWidget(controls_widget)
            
            # Add video container to the splitter
            self.content_splitter.addWidget(video_container)
            
            # Set up the webview with the YouTube player
            success, callback = setup_youtube_webview(
                self.webview, 
                self.document, 
                video_id, 
                target_position,
                self.db_session
            )
            
            if success and callback:
                # Store references for later use
                self.youtube_callback = callback
                self.web_view = self.webview
            
                
                # Check for transcript in metadata
                if hasattr(self.document, 'file_path') and self.document.file_path and os.path.exists(self.document.file_path):
                    try:
                        with open(self.document.file_path, 'r', encoding='utf-8') as f:
                            metadata = json.load(f)
                            if 'transcript' in metadata and metadata['transcript']:
                                # Create transcript view
                                self.transcript_view = YouTubeTranscriptView(
                                    self.db_session,
                                    document_id=self.document_id,
                                    metadata_file=self.document.file_path
                                )
                                self.transcript_view.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
                                
                                # Connect extract created signal
                                self.transcript_view.extractCreated.connect(self.extractCreated)
                                
                                # Add to splitter
                                self.content_splitter.addWidget(self.transcript_view)
                                
                                # Set initial sizes (70% video, 30% transcript)
                                self.content_splitter.setSizes([700, 300])
                                
                                # Connect transcript seek signals
                                self.transcript_view.seek_to_time.connect(self._on_seek_youtube_position)
                            else:
                                # No transcript, show message
                                no_transcript_widget = QWidget()
                                no_transcript_layout = QVBoxLayout(no_transcript_widget)
                                no_transcript_layout.setContentsMargins(10, 10, 10, 10)  # Smaller margins
                                no_transcript_layout.setSpacing(5)  # Reduce spacing between widgets
                                
                                # Create a compact message layout
                                message_widget = QWidget()
                                message_layout = QHBoxLayout(message_widget)
                                message_layout.setContentsMargins(5, 5, 5, 5)
                                
                                # Create the reimport button
                                reimport_button = QPushButton("Reimport Video with Transcript")
                                reimport_button.clicked.connect(self._on_reimport_youtube)
                                reimport_button.setMinimumWidth(200)
                                reimport_button.setMaximumWidth(250)
                                
                                # Create the info label with more compact text
                                no_transcript_label = QLabel(
                                    "No transcript available. Possible reasons: "
                                    "• Disabled by creator • No captions "
                                    "• Age-restricted video • Private content"
                                )
                                no_transcript_label.setWordWrap(True)
                                no_transcript_label.setStyleSheet("color: #555; font-size: 11px; background-color: #f7f7f7; padding: 8px; border-radius: 3px;")
                                
                                # Add widgets to message layout
                                message_layout.addWidget(reimport_button)
                                message_layout.addWidget(no_transcript_label, 1)  # Give label stretch priority
                                
                                # Add message widget to main layout
                                no_transcript_layout.addWidget(message_widget)
                                
                                # Add a spacer to push content to the top
                                no_transcript_layout.addStretch(1)
                                
                                # Add to splitter
                                self.content_splitter.addWidget(no_transcript_widget)
                                
                                # Set initial sizes (85% video, 15% message)
                                self.content_splitter.setSizes([850, 150])
                    except Exception as e:
                        logger.warning(f"Could not load transcript metadata: {e}")
                
                # Add the splitter to the main layout and make it expand
                self.content_layout.addWidget(self.content_splitter, stretch=1)
                
                logger.info(f"Loaded YouTube video: {video_id}")
                return True
            else:
                error_msg = f"Failed to set up YouTube player for video {video_id}"
                self.youtube_status.setText(error_msg)
                self.youtube_status.setStyleSheet("color: red; background-color: #fee; padding: 5px; border-radius: 3px;")
                logger.error(error_msg)
                
                # Add the container to the main layout anyway to show the error
                self.content_layout.addWidget(self.content_splitter, stretch=1)
                return False
                
        except Exception as e:
            logger.exception(f"Error loading YouTube content: {e}")
            from PyQt6.QtWidgets import QLabel
            error_widget = QLabel(f"Error loading YouTube video: {str(e)}")
            error_widget.setStyleSheet("color: red; padding: 20px;")
            error_widget.setWordWrap(True)
            self.content_layout.addWidget(error_widget)
            return False
    
    def _center_widget(self, widget):
        """Helper method to center a widget horizontally."""
        hbox = QHBoxLayout()
        hbox.addStretch(1)
        hbox.addWidget(widget)
        hbox.addStretch(1)
        return hbox

    def _clear_content_layout(self):
        """Clear all widgets from the content layout."""
        if hasattr(self, 'content_layout'):
            # Remove all widgets from the layout
            while self.content_layout.count():
                item = self.content_layout.takeAt(0)
                widget = item.widget()
                if widget:
                    widget.setParent(None)
                    widget.deleteLater()

    def _handle_webview_selection(self, selected_text):
        """Handle text selection from web view."""
        if selected_text and selected_text.strip():
            self.selected_text = selected_text.strip()
    
    @pyqtSlot(QPoint)
    def _on_content_menu(self, pos):
        """Show context menu for document content."""
        # Create menu
        menu = QMenu(self)
        
        # Add actions
        if hasattr(self, 'selected_text') and self.selected_text:
            create_extract_action = menu.addAction("Create Extract")
            create_extract_action.triggered.connect(self._on_create_extract)
        
        # Show menu
        menu.exec(self.mapToGlobal(pos))
    
    @pyqtSlot()
    def _on_previous(self):
        """Navigate to previous page/section."""
        # Implementation depends on document type
        pass
    
    @pyqtSlot()
    def _on_next(self):
        """Navigate to next page/section."""
        # Implementation depends on document type
        pass
    
    @pyqtSlot(int)
    def _on_extract_selected(self, extract_id):
        """Handle extract selection from extract view."""
        self.extractCreated.emit(extract_id)
    
    @pyqtSlot()
    def _on_create_extract(self):
        """Create an extract from selected text."""
        if not self.selected_text:
            return
        
        # Different handling based on the type of content editor
        if HAS_WEBENGINE and isinstance(self.content_edit, QWebEngineView):
            # For QWebEngineView, we already have the selected text from _handle_webview_selection
            # We need to get context differently since we don't have textCursor
            
            # Special handling for YouTube videos
            if hasattr(self, 'document') and self.document and self.document.content_type == 'youtube':
                # For YouTube, use the content_text which should have video info
                context = self.content_text
                position = "youtube"
                if hasattr(self, 'youtube_callback') and self.youtube_callback:
                    # Include current position in the video if available
                    if hasattr(self.youtube_callback, 'current_position'):
                        position = f"youtube:{self.youtube_callback.current_position}"
            # Regular WebView (EPUB, HTML)
            elif self.content_text:
                # Try to find the selection in the content_text to get proper context
                selection_index = self.content_text.find(self.selected_text)
                if selection_index >= 0:
                    # Get surrounding context
                    start_pos = max(0, selection_index - 100)
                    end_pos = min(len(self.content_text), selection_index + len(self.selected_text) + 100)
                    context = self.content_text[start_pos:end_pos]
                    position = f"pos:{selection_index}"
                else:
                    # Fallback if we can't find the exact text
                    context = self.selected_text
                    position = "unknown"
            else:
                context = self.selected_text
                position = "unknown"
        else:
            # For QTextEdit and similar widgets, use the textCursor to get context
            cursor = self.content_edit.textCursor()
            position = cursor.position()
            
            # Try to get some context before and after selection
            start_pos = max(0, position - 100)
            end_pos = min(len(self.content_text), position + 100)
            
            context = self.content_text[start_pos:end_pos]
            position = f"pos:{position}"
        
        # Create the extract
        extract = Extract(
            content=self.selected_text,
            context=context,
            document_id=self.document_id,
            position=position,
            created_date=datetime.utcnow()
        )
        
        try:
            # Add to database
            self.db_session.add(extract)
            self.db_session.commit()
            
            # Refresh extracts view
            self.extract_view.load_extracts_for_document(self.document_id)
            
            # Emit signal
            self.extractCreated.emit(extract.id)
            
            # Clear selection
            self.selected_text = ""
            
        except Exception as e:
            logger.exception(f"Error creating extract: {e}")
            QMessageBox.warning(
                self, "Error", 
                f"Failed to create extract: {str(e)}"
            )
    
    def showEvent(self, event):
        """Handle widget show event to restore view state."""
        super().showEvent(event)
        
        # Restore view state when tab is shown again
        self._restore_view_state()
        
        # Emit signal when the document is shown
        QTimer.singleShot(100, self._on_tab_activated)
        
    def hideEvent(self, event):
        """Handle widget hide event to save view state."""
        super().hideEvent(event)
        
        # Save view state when tab is hidden
        self._save_view_state()
        
        # Emit signal when the document is hidden
        QTimer.singleShot(0, self._on_tab_deactivated)
    
    def _on_tab_activated(self):
        """Handle tab activation."""
        # Additional things to do when a tab becomes active
        # For example, update toolbar actions or refresh content
        logger.debug(f"Tab activated for document: {self.document_id}")
        
        # Special handling for different content types
        if hasattr(self, 'document') and self.document:
            # For PDF content, ensure proper restoration
            if self.document.content_type == 'pdf' and hasattr(self, 'content_edit'):
                # If it's a PDFViewWidget, call its specific methods
                if hasattr(self.content_edit, 'set_view_state') and hasattr(self.content_edit, 'get_view_state'):
                    # Any specific PDF restoration
                    pass
            
            # For YouTube content, which uses web_view instead of content_edit
            elif self.document.content_type == 'youtube' and hasattr(self, 'web_view'):
                logger.debug(f"Tab activated for YouTube video: {self.document.id}")
                # YouTube-specific restoration if needed
                pass
                    
            # For general web content, ensure JavaScript is running
            elif hasattr(self, 'content_edit') and HAS_WEBENGINE and isinstance(self.content_edit, QWebEngineView):
                # Refresh web content to ensure JavaScript is working
                refresh_script = """
                if (typeof refreshActiveContent === 'function') {
                    refreshActiveContent();
                }
                """
                self.content_edit.page().runJavaScript(refresh_script)
    
    def _on_tab_deactivated(self):
        """Handle tab deactivation."""
        # Additional things to do when a tab becomes inactive
        logger.debug(f"Tab deactivated for document: {self.document_id}")
        
        # Save any unsaved changes or state
        self._save_position()
        
        # Content type specific handling
        if hasattr(self, 'document') and self.document:
            # PDF-specific handling
            if self.document.content_type == 'pdf' and hasattr(self, 'content_edit'):
                # Additional PDF-specific state saving
                pass
            
            # YouTube-specific handling
            elif self.document.content_type == 'youtube' and hasattr(self, 'web_view'):
                # Save YouTube position if needed
                # This is usually handled automatically by the setup_youtube_webview callback
                pass
    
    def _save_view_state(self):
        """Save the current view state for later restoration."""
        try:
            # Handle YouTube specially since it uses web_view instead of content_edit
            if hasattr(self, 'document') and self.document and self.document.content_type == 'youtube':
                if hasattr(self, 'web_view') and HAS_WEBENGINE:
                    # Using JavaScript to get YouTube player state
                    self.web_view.page().runJavaScript(
                        "getPlayerState();",
                        lambda state: setattr(self, 'view_state', {**self.view_state, "youtube_state": state}) if state else None
                    )
                return
            
            # If content_edit is a PDFViewWidget, use its specific methods
            if hasattr(self, 'content_edit'):
                # Special handling for PDF view widget
                if hasattr(self.content_edit, 'get_view_state'):
                    pdf_state = self.content_edit.get_view_state()
                    self.view_state.update(pdf_state)
                    logger.debug(f"Saved PDF-specific view state: {pdf_state}")
                elif hasattr(self.content_edit, 'zoom_factor'):
                    self.view_state["zoom_factor"] = self.content_edit.zoom_factor
                
                # Save scroll position for QTextEdit and similar
                if hasattr(self.content_edit, 'verticalScrollBar'):
                    scrollbar = self.content_edit.verticalScrollBar()
                    if scrollbar:
                        self.view_state["scroll_position"] = scrollbar.value()
                        
                # Save scroll position for webviews
                if HAS_WEBENGINE and isinstance(self.content_edit, QWebEngineView):
                    # Using JavaScript to get scroll position
                    self.content_edit.page().runJavaScript(
                        "window.scrollY || document.documentElement.scrollTop || document.body.scrollTop || 0;",
                        lambda pos: setattr(self, 'view_state', {**self.view_state, "scroll_position": pos})
                    )
                    
            # Save size and position if needed
            self.view_state["size"] = self.size()
            
            # Store state in database if needed for persistent storage across sessions
            if hasattr(self, 'document') and self.document:
                # In a production version, you might want to store this in the database
                # import json
                # self.document.view_state = json.dumps(self.view_state)
                # self.db_session.commit()
                pass
                
            logger.debug(f"Saved view state for document {self.document_id}: {self.view_state}")
                
        except Exception as e:
            logger.exception(f"Error saving view state: {e}")
    
    def _restore_view_state(self):
        """Restore the previously saved view state."""
        try:
            if not hasattr(self, 'content_edit') or not self.content_edit:
                return
                
            # If content_edit is a PDFViewWidget, use its specific methods
            if hasattr(self.content_edit, 'set_view_state'):
                # Create a copy of the view state to avoid modifying the original
                pdf_state = {k: v for k, v in self.view_state.items() 
                            if k in ['page', 'zoom_factor', 'position']}
                
                if pdf_state:
                    self.content_edit.set_view_state(pdf_state)
                    logger.debug(f"Restored PDF-specific view state: {pdf_state}")
                return  # Exit early since PDF view handles its own state
                
            # For other view types, apply generic restoration
            
            # Restore zoom factor
            if "zoom_factor" in self.view_state and self.view_state["zoom_factor"]:
                if hasattr(self.content_edit, 'set_zoom'):
                    self.content_edit.set_zoom(self.view_state["zoom_factor"])
                
            # Restore scroll position
            if "scroll_position" in self.view_state and self.view_state["scroll_position"] is not None:
                # For QTextEdit and similar
                if hasattr(self.content_edit, 'verticalScrollBar'):
                    scrollbar = self.content_edit.verticalScrollBar()
                    if scrollbar:
                        scrollbar.setValue(self.view_state["scroll_position"])
                        
                # For web views
                if HAS_WEBENGINE and isinstance(self.content_edit, QWebEngineView):
                    pos = self.view_state["scroll_position"]
                    script = f"window.scrollTo(0, {pos});"
                    self.content_edit.page().runJavaScript(script)
                    
            # Apply sizing if needed
            if "size" in self.view_state and self.view_state["size"]:
                # Usually not needed as the tab widget will control size,
                # but could be useful in some cases
                pass
                
            logger.debug(f"Restored view state for document {self.document_id}: {self.view_state}")
                
        except Exception as e:
            logger.exception(f"Error restoring view state: {e}")
    
    def closeEvent(self, event):
        try:
            # Force save YouTube position if applicable
            if (hasattr(self, 'document') and self.document and 
                self.document.content_type == 'youtube' and 
                hasattr(self, 'youtube_callback') and self.youtube_callback):
                
                logger.info(f"Saving YouTube position on close: {self.youtube_callback.current_position}")
                self.youtube_callback.savePosition()
                
            self._save_position()
            self._save_view_state()
            self.db_session.commit()  # Ensure changes are committed
        except Exception as e:
            logger.exception(f"Error in closeEvent: {e}")
        super().closeEvent(event)
    
    def _save_position(self):
        """Save the current reading position."""
        try:
            # Check if we have a document to save position for
            if not hasattr(self, 'document') or not self.document:
                return
            
            # For QTextEdit
            if hasattr(self, 'text_view') and isinstance(self.text_view, QTextEdit):
                cursor = self.text_view.textCursor()
                position = cursor.position()
                
                # Update the document
                self.document.position = position
                self.db_session.commit()
                logger.debug(f"Saved text cursor position: {position}")
            
            # For QWebEngineView (position handled by the helper module)
            elif hasattr(self, 'web_view') and QWebEngineView and isinstance(self.web_view, QWebEngineView):
                # Position is saved automatically by the setup_youtube_webview callback system
                logger.debug("WebView position handled by helper module")
                
        except Exception as e:
            logger.exception(f"Error saving position: {e}")
    
    def _restore_position(self):
        """Restore the last reading position."""
        try:
            if not hasattr(self, 'document') or not self.document:
                return
                
            # Get stored position
            position = getattr(self.document, 'position', None)
            if position is None or position <= 0:
                logger.info(f"No stored position found for {self.document.title}")
                return
                    
            logger.info(f"Attempting to restore position {position} for {self.document.title}")
                
            # Determine how to set position based on document type and view
            if hasattr(self, 'content_edit'):
                if isinstance(self.content_edit, QTextEdit):
                    # For text documents, set cursor position
                    cursor = self.content_edit.textCursor()
                        
                    # Make sure position is within valid range
                    doc_length = len(self.content_edit.toPlainText())
                    position = min(position, doc_length)
                        
                    cursor.setPosition(position)
                    self.content_edit.setTextCursor(cursor)
                        
                    # Ensure the cursor is visible
                    self.content_edit.ensureCursorVisible()
                    logger.info(f"Restored text document position: {position}")
                        
                elif HAS_WEBENGINE and isinstance(self.content_edit, QWebEngineView):
                    # For WebEngine view (EPUB/HTML/YouTube), position is handled by the helper module
                    # No need to do anything here as it's done when the view is initialized
                    pass
                else:
                    logger.warning(f"Unknown content editor type, can't restore position: {type(self.content_edit)}")
            else:
                logger.warning("No content_edit available to restore position to")
                
        except Exception as e:
            logger.exception(f"Error restoring document position: {e}")

    def _on_reimport_youtube(self):
        """Reimport the YouTube video with transcript."""
        try:
            from PyQt6.QtWidgets import QMessageBox
            
            # Show confirmation dialog
            reply = QMessageBox.question(
                self, 
                "Reimport YouTube Video",
                "This will reimport the YouTube video and attempt to fetch the transcript again. Continue?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                QMessageBox.StandardButton.Yes
            )
            
            if reply == QMessageBox.StandardButton.Yes:
                # Check if we have a source URL or video ID
                if hasattr(self.document, 'source_url') and self.document.source_url:
                    source_url = self.document.source_url
                    
                    # Import the document processor in the handler to avoid circular imports
                    from core.document_processor.handlers.youtube_handler import YouTubeHandler
                    handler = YouTubeHandler()
                    
                    # Extract video ID from source URL
                    video_id = handler._extract_video_id(source_url)
                    
                    if video_id:
                        # Show importing message
                        self.youtube_status.setText(f"Reimporting video {video_id} with transcript...")
                        self.youtube_status.setStyleSheet("color: #000; background-color: #ffd; padding: 5px; border-radius: 3px;")
                        
                        # Process in a background thread to avoid UI freezing
                        from PyQt6.QtCore import QThread, pyqtSignal
                        
                        class ImportThread(QThread):
                            importFinished = pyqtSignal(bool, str)
                            
                            def __init__(self, handler, url, parent=None):
                                super().__init__(parent)
                                self.handler = handler
                                self.url = url
                                
                            def run(self):
                                try:
                                    # Reimport with force_transcript=True
                                    success = self.handler.process_url(self.url, force_transcript=True)
                                    self.importFinished.emit(success, "")
                                except Exception as e:
                                    self.importFinished.emit(False, str(e))
                        
                        # Create and start the thread
                        self.import_thread = ImportThread(handler, source_url, self)
                        self.import_thread.importFinished.connect(self._on_reimport_finished)
                        self.import_thread.start()
                    else:
                        self.youtube_status.setText("Could not extract video ID from URL")
                        self.youtube_status.setStyleSheet("color: red; background-color: #fee; padding: 5px; border-radius: 3px;")
                else:
                    self.youtube_status.setText("No source URL available for reimport")
                    self.youtube_status.setStyleSheet("color: red; background-color: #fee; padding: 5px; border-radius: 3px;")
        except Exception as e:
            logger.exception(f"Error starting YouTube reimport: {e}")
            self.youtube_status.setText(f"Error: {str(e)}")
            self.youtube_status.setStyleSheet("color: red; background-color: #fee; padding: 5px; border-radius: 3px;")
            
    def _on_reimport_finished(self, success, error_msg):
        """Handle completion of YouTube reimport."""
        if success:
            self.youtube_status.setText("Video reimported successfully. Reloading...")
            self.youtube_status.setStyleSheet("color: green; background-color: #efe; padding: 5px; border-radius: 3px;")
            
            # Reload the document after a short delay
            from PyQt6.QtCore import QTimer
            QTimer.singleShot(1500, lambda: self.load_document(self.document_id))
        else:
            self.youtube_status.setText(f"Reimport failed: {error_msg}")
            self.youtube_status.setStyleSheet("color: red; background-color: #fee; padding: 5px; border-radius: 3px;")

    def save_document(self):
        """Save the current document state to the database."""
        try:
            if not self.document or not self.document_id:
                logger.debug("No document to save")
                return False
                
            # Save position
            self._save_position()
            
            # Save view state
            self._save_view_state()
            
            # Update last_modified timestamp
            self.document.last_modified = datetime.utcnow()
            
            # Commit changes to the database
            self.db_session.commit()
            
            logger.debug(f"Document {self.document_id} saved successfully")
            return True
            
        except Exception as e:
            logger.exception(f"Error saving document: {e}")
            return False

    def _update_youtube_position(self):
        """Update and save the current YouTube position."""
        if not hasattr(self, 'web_view') or not self.web_view:
            return
            
        # Use JavaScript to get the current position
        self.web_view.page().runJavaScript(
            "getCurrentPosition();",
            self._handle_position_update
        )
        
    def _handle_position_update(self, position):
        """Handle position update from JavaScript."""
        if position is None or not isinstance(position, (int, float)):
            return
            
        # Update position label
        if hasattr(self, 'position_label'):
            self.position_label.setText(f"Position: {int(position)}s")
        
        # Update the document position in the database
        if hasattr(self, 'document') and self.document:
            try:
                if position > 0:  # Don't save if at the beginning
                    self.document.position = int(position)
                    self.db_session.commit()
                    logger.debug(f"Saved YouTube position: {position}")
            except Exception as e:
                logger.error(f"Error saving YouTube position: {e}")
                
    def _on_save_youtube_position(self):
        """Manually save the current YouTube position."""
        if hasattr(self, 'youtube_callback') and self.youtube_callback:
            try:
                self.youtube_callback.savePosition()
                logger.debug(f"Position saved: {self.youtube_callback.current_position}s")
            except Exception as e:
                logger.error(f"Error manually saving position: {e}")
        else:
            self._update_youtube_position()  # Fallback
        
    def _on_seek_youtube_position(self):
        """Handle seeking to a specific position in a YouTube video."""
        if not hasattr(self, 'web_view') or not self.web_view or not hasattr(self, 'seek_time_input'):
            return
            
        try:
            # Get position from input
            time_text = self.seek_time_input.text()
            position = int(time_text)
            
            # Use JavaScript to seek
            seek_script = f"seekToTime({position});"
            self.web_view.page().runJavaScript(seek_script)
            
            # Update backend
            if hasattr(self, 'youtube_callback') and self.youtube_callback:
                self.youtube_callback.current_position = position
                self.youtube_callback.onTimeUpdate(position)
                
            # Update label
            if hasattr(self, 'position_label'):
                self.position_label.setText(f"Position: {position}s")
                
        except (ValueError, TypeError) as e:
            logger.error(f"Invalid position value: {e}")

    def _set_error_message(self, message):
        """
        Display an error message in the content area.
        
        Args:
            message: Error message to display
        """
        # Clear the content layout
        self._clear_content_layout()
        
        # Create error label
        error_widget = QLabel(message)
        error_widget.setWordWrap(True)
        error_widget.setStyleSheet("color: red; padding: 20px;")
        
        # Add to layout
        self.content_layout.addWidget(error_widget)
        
        logger.error(message)

    def _add_supermemo_features(self, web_view):
        """
        Add SuperMemo SM-18 style incremental reading features to web view.
        
        Args:
            web_view: QWebEngineView instance
        """
        if not HAS_WEBENGINE or not isinstance(web_view, QWebEngineView):
            return
            
        logger.info("Adding SuperMemo incremental reading features to EPUB view")
        
        # Add SuperMemo toolbar
        toolbar = QToolBar("SuperMemo Features")
        toolbar.setObjectName("superMemoToolbar")
        
        # Extract action
        extract_action = QAction("Extract Selection", self)
        extract_action.setToolTip("Extract selected text to create a new learning item (Ctrl+Alt+E)")
        extract_action.setShortcut("Ctrl+Alt+E")
        extract_action.triggered.connect(self._on_sm_extract)
        
        # Cloze deletion action
        cloze_action = QAction("Cloze Deletion", self)
        cloze_action.setToolTip("Create cloze deletion from selected text (Ctrl+Alt+C)")
        cloze_action.setShortcut("Ctrl+Alt+C")
        cloze_action.triggered.connect(self._on_sm_cloze)
        
        # Priority tools
        priority_action = QAction("Set Priority", self)
        priority_action.setToolTip("Set learning priority for this item (Ctrl+Alt+P)")
        priority_action.setShortcut("Ctrl+Alt+P")
        priority_action.triggered.connect(self._on_sm_priority)
        
        # Highlighting tools with SuperMemo color scheme
        highlight_menu = QMenu("Highlight", self)
        
        # SM-18 style color highlighting
        colors = [
            ("Most Important", "#FF0000", "Ctrl+Alt+1"),  # Red
            ("Very Important", "#FF7F00", "Ctrl+Alt+2"),  # Orange
            ("Important", "#FFFF00", "Ctrl+Alt+3"),       # Yellow
            ("Less Important", "#00FF00", "Ctrl+Alt+4"),  # Green
            ("Least Important", "#0000FF", "Ctrl+Alt+5")  # Blue
        ]
        
        for name, color, shortcut in colors:
            action = QAction(name, self)
            action.setData(color)
            action.setShortcut(shortcut)
            action.triggered.connect(lambda checked, c=color: self._on_sm_highlight(c))
            highlight_menu.addAction(action)
        
        # Schedule review actions
        schedule_action = QAction("Schedule Review", self)
        schedule_action.setToolTip("Schedule this item for review (Ctrl+Alt+S)")
        schedule_action.setShortcut("Ctrl+Alt+S")
        schedule_action.triggered.connect(self._on_sm_schedule)
        
        # Add actions to toolbar
        toolbar.addAction(extract_action)
        toolbar.addAction(cloze_action)
        toolbar.addAction(priority_action)
        toolbar.addSeparator()
        toolbar.addAction(schedule_action)
        toolbar.addSeparator()
        
        # Add highlight menu button
        highlight_button = QPushButton("Highlight")
        highlight_button.setMenu(highlight_menu)
        toolbar.addWidget(highlight_button)
        
        # Add spacing
        spacer = QWidget()
        spacer.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
        toolbar.addWidget(spacer)
        
        # Add toolbar to layout before the web view
        if hasattr(self, 'content_layout'):
            # Insert the toolbar at the top
            self.content_layout.insertWidget(0, toolbar)
        
        # Add JavaScript for SuperMemo features
        js_code = """
        (function() {
            // SuperMemo Incremental Reading Enhancements
            document.addEventListener('keydown', function(e) {
                // Ctrl+Alt+Z - Fast forward (skip current item)
                if (e.ctrlKey && e.altKey && e.code === 'KeyZ') {
                    window.superMemo.skipItem();
                }
            });
            
            // Make text easier to process in chunks
            const style = document.createElement('style');
            style.textContent = `
                p, li { 
                    margin-bottom: 1em; 
                    line-height: 1.5;
                }
                .sm-highlight-red { background-color: rgba(255,0,0,0.3); }
                .sm-highlight-orange { background-color: rgba(255,127,0,0.3); }
                .sm-highlight-yellow { background-color: rgba(255,255,0,0.3); }
                .sm-highlight-green { background-color: rgba(0,255,0,0.3); }
                .sm-highlight-blue { background-color: rgba(0,0,255,0.3); }
                .sm-extracted { border-left: 3px solid #0078D7; padding-left: 10px; }
            `;
            document.head.appendChild(style);
            
            window.superMemo = {
                highlight: function(color) {
                    const selection = window.getSelection();
                    if (selection.rangeCount > 0) {
                        const range = selection.getRangeAt(0);
                        const span = document.createElement('span');
                        span.className = 'sm-highlight-' + color;
                        range.surroundContents(span);
                    }
                },
                
                extract: function() {
                    const selection = window.getSelection();
                    if (selection.rangeCount > 0) {
                        const range = selection.getRangeAt(0);
                        const div = document.createElement('div');
                        div.className = 'sm-extracted';
                        range.surroundContents(div);
                        return selection.toString();
                    }
                    return '';
                },
                
                skipItem: function() {
                    // This would be connected to the review system
                    console.log('Item skipped for later review');
                }
            };
        })();
        """
        
        # Inject the SuperMemo JavaScript
        web_view.page().runJavaScript(js_code)
        
        # Connect JavaScript bridge for communication
        self.sm_channel = QWebChannel()
        self.sm_callback = WebViewCallback(self)
        self.sm_channel.registerObject("superMemoHandler", self.sm_callback)
        web_view.page().setWebChannel(self.sm_channel)
        
        # Keep references alive
        self.keep_alive(toolbar)
        self.keep_alive(self.sm_channel)
        self.keep_alive(self.sm_callback)
    
    @pyqtSlot()
    def _on_sm_extract(self):
        """Create a SuperMemo-style extract from selected text."""
        if hasattr(self, 'content_edit') and isinstance(self.content_edit, QWebEngineView):
            # Get selected text via JavaScript
            self.content_edit.page().runJavaScript(
                "window.superMemo.extract();",
                self._handle_sm_extract_result
            )
    
    def _handle_sm_extract_result(self, selected_text):
        """Handle the extract creation from selected text."""
        if not selected_text:
            return
            
        try:
            # Create a new extract with SuperMemo metadata
            from core.knowledge_base.models import Extract
            
            # Include SuperMemo priority and scheduling data
            extract = Extract(
                document_id=self.document.id,
                content=selected_text,
                source_page=getattr(self.document, 'current_page', 0),
                created_at=datetime.now(),
                metadata={
                    "sm_priority": 50,  # Default priority (0-100)
                    "sm_interval": 1,   # Days until next review
                    "sm_ease_factor": 2.5,  # SuperMemo ease factor
                    "sm_extract_type": "topic"  # SuperMemo concept
                }
            )
            
            # Save to database
            self.db_session.add(extract)
            self.db_session.commit()
            
            # Show confirmation
            QApplication.instance().beep()
            
            # Notify about the new extract
            if hasattr(self, 'extractCreated'):
                self.extractCreated.emit(extract.id)
                
            logger.info(f"Created SuperMemo extract {extract.id} with priority 50")
            
        except Exception as e:
            logger.exception(f"Error creating SuperMemo extract: {e}")
    
    @pyqtSlot()
    def _on_sm_cloze(self):
        """Create a SuperMemo-style cloze deletion from selected text."""
        if hasattr(self, 'content_edit') and isinstance(self.content_edit, QWebEngineView):
            # Get selected text via JavaScript
            self.content_edit.page().runJavaScript(
                "window.getSelection().toString();",
                self._handle_sm_cloze_result
            )
    
    def _handle_sm_cloze_result(self, selected_text):
        """Handle the cloze deletion creation from selected text."""
        if not selected_text:
            return
            
        try:
            # Create a new learning item with cloze deletion
            from core.knowledge_base.models import LearningItem
            
            # Create cloze format: text with [...] replacing the selection
            content = self.content_text
            cloze_text = content.replace(selected_text, f"[...]")
            
            # Create the learning item with SuperMemo properties
            item = LearningItem(
                content=cloze_text,
                answer=selected_text,
                item_type="cloze",
                document_id=self.document.id,
                metadata={
                    "sm_priority": 60,  # Default priority for cloze (0-100)
                    "sm_interval": 1,   # Days until next review
                    "sm_ease_factor": 2.5,  # SuperMemo ease factor
                    "sm_repetitions": 0  # Number of reviews
                }
            )
            
            # Save to database
            self.db_session.add(item)
            self.db_session.commit()
            
            # Show confirmation
            QApplication.instance().beep()
            
            logger.info(f"Created SuperMemo cloze item {item.id} with priority 60")
            
        except Exception as e:
            logger.exception(f"Error creating SuperMemo cloze: {e}")
    
    @pyqtSlot()
    def _on_sm_priority(self):
        """Set SuperMemo priority for the current document or selection."""
        # Create a dialog for setting priority
        dialog = QDialog(self)
        dialog.setWindowTitle("Set SuperMemo Priority")
        layout = QVBoxLayout()
        
        # Slider for priority (0-100)
        priority_slider = QSlider(Qt.Orientation.Horizontal)
        priority_slider.setMinimum(0)
        priority_slider.setMaximum(100)
        priority_slider.setValue(50)  # Default
        priority_slider.setTickPosition(QSlider.TickPosition.TicksBelow)
        priority_slider.setTickInterval(10)
        
        # Label to show value
        value_label = QLabel("Priority: 50")
        priority_slider.valueChanged.connect(lambda v: value_label.setText(f"Priority: {v}"))
        
        # Add to layout
        layout.addWidget(QLabel("Set learning priority (0-100):"))
        layout.addWidget(priority_slider)
        layout.addWidget(value_label)
        
        # Buttons
        buttons = QHBoxLayout()
        ok_button = QPushButton("OK")
        cancel_button = QPushButton("Cancel")
        
        ok_button.clicked.connect(dialog.accept)
        cancel_button.clicked.connect(dialog.reject)
        
        buttons.addWidget(ok_button)
        buttons.addWidget(cancel_button)
        
        layout.addLayout(buttons)
        dialog.setLayout(layout)
        
        # Show dialog and process result
        if dialog.exec() == QDialog.DialogCode.Accepted:
            priority = priority_slider.value()
            
            # Update document metadata with SuperMemo priority
            if hasattr(self, 'document') and self.document:
                metadata = self.document.metadata or {}
                metadata['sm_priority'] = priority
                self.document.metadata = metadata
                
                # Save changes
                self.db_session.commit()
                
                logger.info(f"Set SuperMemo priority {priority} for document {self.document.id}")
    
    @pyqtSlot(str)
    def _on_sm_highlight(self, color):
        """Apply SuperMemo-style highlighting to selected text."""
        if hasattr(self, 'content_edit') and isinstance(self.content_edit, QWebEngineView):
            # Map the color to a color name for JavaScript
            color_map = {
                "#FF0000": "red",     # Red
                "#FF7F00": "orange",  # Orange
                "#FFFF00": "yellow",  # Yellow
                "#00FF00": "green",   # Green
                "#0000FF": "blue"     # Blue
            }
            
            color_name = color_map.get(color, "yellow")
            
            # Apply the highlight via JavaScript
            self.content_edit.page().runJavaScript(
                f"window.superMemo.highlight('{color_name}');"
            )
            
            logger.debug(f"Applied SuperMemo {color_name} highlight")
    
    @pyqtSlot()
    def _on_sm_schedule(self):
        """Schedule the current document for SuperMemo-style review."""
        # Create dialog for scheduling
        dialog = QDialog(self)
        dialog.setWindowTitle("Schedule Review")
        layout = QVBoxLayout()
        
        # Days until next review
        days_spinner = QLineEdit()
        days_spinner.setText("1")  # Default: tomorrow
        days_spinner.setValidator(QIntValidator(1, 365))
        
        # Add to layout
        layout.addWidget(QLabel("Days until next review:"))
        layout.addWidget(days_spinner)
        
        # Add SuperMemo Algorithm options
        algorithm_checkbox = QCheckBox("Use SuperMemo SM-18 algorithm")
        algorithm_checkbox.setChecked(True)
        layout.addWidget(algorithm_checkbox)
        
        # Buttons
        buttons = QHBoxLayout()
        ok_button = QPushButton("Schedule")
        cancel_button = QPushButton("Cancel")
        
        ok_button.clicked.connect(dialog.accept)
        cancel_button.clicked.connect(dialog.reject)
        
        buttons.addWidget(ok_button)
        buttons.addWidget(cancel_button)
        
        layout.addLayout(buttons)
        dialog.setLayout(layout)
        
        # Show dialog and process result
        if dialog.exec() == QDialog.DialogCode.Accepted:
            try:
                days = int(days_spinner.text())
                use_sm_algorithm = algorithm_checkbox.isChecked()
                
                # Update document metadata with scheduling information
                if hasattr(self, 'document') and self.document:
                    metadata = self.document.metadata or {}
                    
                    # Calculate next review date
                    from datetime import datetime, timedelta
                    next_review = datetime.now() + timedelta(days=days)
                    
                    # Update metadata
                    metadata['sm_next_review'] = next_review.isoformat()
                    metadata['sm_interval'] = days
                    
                    if use_sm_algorithm:
                        # Apply SuperMemo SM-18 algorithm factors
                        metadata['sm_ease_factor'] = metadata.get('sm_ease_factor', 2.5)
                        metadata['sm_repetitions'] = metadata.get('sm_repetitions', 0) + 1
                    
                    self.document.metadata = metadata
                    self.document.last_reviewed = datetime.now()
                    
                    # Save changes
                    self.db_session.commit()
                    
                    logger.info(f"Scheduled document {self.document.id} for review in {days} days")
                    
                    # Show confirmation
                    QMessageBox.information(
                        self,
                        "Review Scheduled",
                        f"Document scheduled for review in {days} days."
                    )
                    
            except Exception as e:
                logger.exception(f"Error scheduling review: {e}")
                QMessageBox.warning(
                    self,
                    "Error",
                    f"Could not schedule review: {str(e)}"
                )

    # Import necessary module to avoid error
    from datetime import datetime

    def _update_theme(self):
        """Update the theme for this document view."""
        try:
            # Apply theme to the content if it's a web view
            if hasattr(self, 'content_edit') and isinstance(self.content_edit, QWebEngineView):
                # Get the theme manager
                from core.utils.theme_manager import ThemeManager
                from PyQt6.QtWidgets import QApplication
                
                # Get or create theme manager
                app = QApplication.instance()
                main_window = None
                for widget in app.topLevelWidgets():
                    if widget.__class__.__name__ == "MainWindow":
                        main_window = widget
                        break
                
                theme_manager = None
                if main_window and hasattr(main_window, 'theme_manager'):
                    theme_manager = main_window.theme_manager
                else:
                    # Create a new instance if not available
                    from core.settings.settings_manager import SettingsManager
                    settings_manager = SettingsManager()
                    theme_manager = ThemeManager(settings_manager)
                
                # Apply theme to the web content
                from ui.load_epub_helper import apply_theme_background
                apply_theme_background(self.content_edit, theme_manager)
                
                logger.debug("Updated theme for document view")
        except Exception as e:
            logger.exception(f"Error updating theme: {e}")
